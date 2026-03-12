"""
Script para generar el Manual de Usuario del Módulo de Facturación - Geolab
Ejecutar: python3 generar_manual_facturacion.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ============================================================
# CONFIGURACIÓN DE COLORES
# ============================================================
BRAND_BLUE = RGBColor(0, 142, 214)
DARK_BLUE = RGBColor(28, 46, 74)
WHITE = RGBColor(255, 255, 255)
GRAY_TEXT = RGBColor(100, 116, 139)
SUCCESS_GREEN = RGBColor(22, 163, 74)
WARNING_ORANGE = RGBColor(234, 88, 12)
DANGER_RED = RGBColor(220, 38, 38)

doc = Document()

# ============================================================
# ESTILOS BASE
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(51, 65, 85)

for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = DARK_BLUE
    if i == 1:
        h.font.size = Pt(26)
        h.font.bold = True
    elif i == 2:
        h.font.size = Pt(20)
        h.font.bold = True
    elif i == 3:
        h.font.size = Pt(15)
        h.font.bold = True
    elif i == 4:
        h.font.size = Pt(12)
        h.font.bold = True


# ============================================================
# FUNCIONES AUXILIARES
# ============================================================
def colored_p(text, color=None, bold=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def tip_box(title, text, icon="💡"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {icon} {title}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_BLUE
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY_TEXT


def warning_box(text):
    tip_box("Importante", text, icon="⚠️")


def step(number, title, description=""):
    p = doc.add_paragraph()
    run_num = p.add_run(f"  Paso {number}: ")
    run_num.bold = True
    run_num.font.color.rgb = BRAND_BLUE
    run_num.font.size = Pt(11)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    if description:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        run_desc = p2.add_run(description)
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = GRAY_TEXT


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return t


def bullet(text):
    doc.add_paragraph(text, style='List Bullet')


def bullet_bold(title, desc):
    p = doc.add_paragraph()
    run = p.add_run(f"• {title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)


def sub_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(2)


def page_break():
    doc.add_page_break()


# ============================================================
# PORTADA
# ============================================================
for _ in range(5):
    doc.add_paragraph()

colored_p("GEOLAB S.A.S", BRAND_BLUE, bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("Portal de Gestión Técnica", DARK_BLUE, bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

colored_p("MÓDULO DE FACTURACIÓN", DARK_BLUE, bold=True, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("Manual de Usuario", DARK_BLUE, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

colored_p("Guía completa para el registro de servicios,", GRAY_TEXT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("gestión de precios, generación de facturas", GRAY_TEXT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("y control del histórico de facturación", GRAY_TEXT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(5):
    doc.add_paragraph()

colored_p("Versión 1.0  |  Marzo 2026", GRAY_TEXT, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("Desarrollado por Vadom Data Consulting SAS", GRAY_TEXT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================
# TABLA DE CONTENIDO
# ============================================================
doc.add_heading('Tabla de Contenido', level=1)
doc.add_paragraph()

toc = [
    ("1.", "Introducción al Módulo de Facturación"),
    ("2.", "Acceso y Navegación"),
    ("3.", "Panel de Facturación (Dashboard)"),
    ("4.", "Catálogo de Servicios"),
    ("   4.1", "Categorías de Servicio"),
    ("   4.2", "Tipos de Servicio"),
    ("   4.3", "Buscar en el Catálogo"),
    ("5.", "Lista de Precios por Obra"),
    ("   5.1", "¿Qué es la lista de precios?"),
    ("   5.2", "Configurar precios paso a paso"),
    ("   5.3", "Precios y congelamiento"),
    ("6.", "Gestión de Impuestos"),
    ("   6.1", "Crear el IVA"),
    ("   6.2", "Regla especial de Transporte"),
    ("7.", "Registro de Servicios Realizados"),
    ("   7.1", "Proceso paso a paso"),
    ("   7.2", "Agregar múltiples líneas"),
    ("   7.3", "Validaciones y errores comunes"),
    ("8.", "Histórico de Registros"),
    ("   8.1", "Filtros de búsqueda"),
    ("   8.2", "Editar un registro"),
    ("   8.3", "Eliminar un registro"),
    ("9.", "Generación de Facturas"),
    ("   9.1", "Seleccionar obra y periodo"),
    ("   9.2", "Vista previa de la factura"),
    ("   9.3", "Servicios normales vs. transporte"),
    ("   9.4", "Selección de IVA en transporte"),
    ("   9.5", "Confirmar y generar"),
    ("   9.6", "PDF de la factura"),
    ("10.", "Repositorio de Facturas"),
    ("   10.1", "Nivel 1: Clientes"),
    ("   10.2", "Nivel 2: Obras del cliente"),
    ("   10.3", "Nivel 3: Facturas de la obra"),
    ("11.", "Acciones sobre Facturas"),
    ("   11.1", "Marcar como Pagada"),
    ("   11.2", "Anular una Factura"),
    ("   11.3", "Descargar PDF"),
    ("12.", "Flujo Completo: Ejemplo Práctico"),
    ("13.", "Preguntas Frecuentes"),
    ("14.", "Glosario de Términos"),
]

for num, titulo in toc:
    p = doc.add_paragraph()
    if num.strip().endswith('.'):
        run_n = p.add_run(f"{num}  ")
        run_n.bold = True
        run_n.font.color.rgb = BRAND_BLUE
        run_n.font.size = Pt(12)
        run_t = p.add_run(titulo)
        run_t.font.size = Pt(12)
        run_t.bold = True
    else:
        run_n = p.add_run(f"    {num.strip()}  ")
        run_n.font.color.rgb = GRAY_TEXT
        run_n.font.size = Pt(11)
        run_t = p.add_run(titulo)
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(51, 65, 85)

page_break()


# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
doc.add_heading('1. Introducción al Módulo de Facturación', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'El Módulo de Facturación del Portal Geolab permite llevar el control completo de los '
    'servicios prestados a cada obra, desde el registro diario hasta la generación de facturas '
    'formales con cálculo automático de impuestos.'
)

doc.add_heading('¿Para qué sirve?', level=3)

doc.add_paragraph(
    'Este módulo resuelve la necesidad de:'
)

beneficios = [
    ("Registrar los servicios realizados", "Cada vez que Geolab realiza un ensayo o servicio para una obra, se registra con fecha, cantidad, precio y referencia al informe."),
    ("Controlar precios por obra", "Cada obra puede tener precios diferentes para los mismos servicios, y el sistema los gestiona de forma independiente."),
    ("Generar facturas automáticas", "A partir de los registros pendientes, el sistema genera la factura con cálculo de IVA, desglose por servicio y PDF profesional."),
    ("Llevar el histórico completo", "Todas las facturas quedan almacenadas en un repositorio organizado por cliente y obra."),
    ("Evitar errores de cálculo", "Los precios se congelan al momento del registro, el IVA se calcula automáticamente y los totales son exactos."),
]
for titulo, desc in beneficios:
    bullet_bold(titulo, desc)

doc.add_paragraph()

doc.add_heading('Flujo general', level=3)
doc.add_paragraph(
    'El proceso de facturación sigue un flujo claro de 5 pasos:'
)

flujo = [
    ("Configurar", "Crear el catálogo de servicios, asignar precios a cada obra y configurar impuestos."),
    ("Registrar", "Cada día registrar los servicios realizados para cada obra."),
    ("Acumular", "Los registros se van acumulando como 'pendientes de facturar'."),
    ("Facturar", "Periódicamente, generar la factura seleccionando obra, periodo y registros."),
    ("Gestionar", "Marcar facturas como pagadas, anular si es necesario, consultar el histórico."),
]
for i, (titulo, desc) in enumerate(flujo, 1):
    p = doc.add_paragraph()
    run_n = p.add_run(f"  {i}. {titulo}  →  ")
    run_n.bold = True
    run_n.font.color.rgb = BRAND_BLUE
    run2 = p.add_run(desc)

doc.add_paragraph()
warning_box("Solo los usuarios con rol Staff Administrativo (Administración, Laboratorio o Recepción) tienen acceso al módulo de facturación. Los técnicos de laboratorio y los clientes no ven estas opciones.")

page_break()


# ============================================================
# 2. ACCESO Y NAVEGACIÓN
# ============================================================
doc.add_heading('2. Acceso y Navegación', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'Todas las funciones de facturación se encuentran agrupadas en la sección FACTURACIÓN '
    'del menú lateral izquierdo (sidebar). Al iniciar sesión como Staff Administrativo, '
    'verá las siguientes opciones:'
)

doc.add_paragraph()

table(
    ["Opción del Menú", "Ícono", "¿Qué hace?"],
    [
        ["Panel Facturación", "📊 Recibo", "Dashboard con estadísticas y gráficos del mes"],
        ["Registrar Servicio", "➕ Más", "Formulario para registrar servicios realizados a una obra"],
        ["Generar Factura", "📄 Documento", "Genera factura a partir de registros pendientes de una obra"],
        ["Histórico", "🕐 Reloj", "Lista filtrable de todos los registros de servicios"],
        ["Repositorio", "📦 Archivo", "Navegación por clientes → obras → facturas"],
        ["Catálogo Servicios", "📋 Lista", "Administración de categorías y tipos de servicio"],
        ["Impuestos", "% Porcentaje", "Configuración de impuestos (IVA)"],
    ]
)

doc.add_paragraph()

doc.add_heading('Acceso rápido desde el expediente de obra', level=3)
doc.add_paragraph(
    'Además del menú lateral, puede acceder a la lista de precios de una obra directamente '
    'desde su expediente. En el encabezado del expediente de obra verá un botón con el '
    'ícono "$" que lo lleva a la gestión de precios de esa obra.'
)

doc.add_paragraph()
doc.add_heading('Buscadores inteligentes', level=3)
doc.add_paragraph(
    'Los campos de selección de Constructora, Obra, y Tipo de Servicio en todo el módulo '
    'de facturación funcionan como buscadores: puede escribir para filtrar las opciones en '
    'lugar de recorrer listas largas. Simplemente empiece a escribir el nombre y las opciones '
    'se filtran automáticamente.'
)

page_break()


# ============================================================
# 3. PANEL DE FACTURACIÓN
# ============================================================
doc.add_heading('3. Panel de Facturación (Dashboard)', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Panel Facturación", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El Panel de Facturación es la pantalla de inicio del módulo. Ofrece una visión general '
    'del estado de la facturación con indicadores en tiempo real.'
)

doc.add_heading('3.1 Indicadores Principales (KPIs)', level=2)
doc.add_paragraph(
    'En la parte superior encontrará tarjetas con los indicadores clave:'
)

kpis = [
    ("Constructoras", "Cantidad total de empresas constructoras registradas en el sistema."),
    ("Obras", "Número total de proyectos/obras activas."),
    ("Facturas este Mes", "Cantidad de facturas generadas en el mes actual y el monto total facturado."),
    ("Registros este Mes", "Cantidad de servicios registrados en el mes actual."),
    ("Registros Pendientes", "Servicios que han sido registrados pero aún no se han incluido en ninguna factura. Este número indica cuánto hay por facturar."),
]
for titulo, desc in kpis:
    bullet_bold(titulo, desc)

doc.add_heading('3.2 Gráfico de Facturación Mensual', level=2)
doc.add_paragraph(
    'Un gráfico de barras muestra la evolución de la facturación durante los últimos 6 meses. '
    'Cada barra representa el monto total facturado en ese mes. Esto permite identificar '
    'tendencias y estacionalidad en la operación.'
)

doc.add_heading('3.3 Top Clientes y Servicios', level=2)
doc.add_paragraph(
    'En la parte inferior del dashboard se muestran dos listados:'
)

bullet_bold("Top 3 Clientes", "Las empresas constructoras con mayor facturación acumulada históricamente. Útil para identificar los clientes más importantes.")
bullet_bold("Top 3 Servicios", "Los tipos de servicio más realizados en los últimos 90 días. Útil para conocer la demanda actual del laboratorio.")

doc.add_paragraph()
tip_box("Actualización", "Los datos del dashboard se actualizan cada vez que ingresa a la pantalla. No necesita recargar manualmente.")

page_break()


# ============================================================
# 4. CATÁLOGO DE SERVICIOS
# ============================================================
doc.add_heading('4. Catálogo de Servicios', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Catálogo Servicios", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El catálogo es el directorio maestro de todos los servicios que ofrece Geolab. '
    'Antes de poder registrar servicios o generar facturas, los servicios deben existir '
    'en el catálogo. Se organiza en dos niveles: Categorías y Tipos de Servicio.'
)

doc.add_heading('4.1 Categorías de Servicio', level=2)
doc.add_paragraph(
    'Las categorías agrupan servicios similares. Cada categoría tiene un código numérico '
    'y un nombre descriptivo.'
)

doc.add_heading('Categorías típicas:', level=3)
table(
    ["Código", "Categoría", "Descripción"],
    [
        ["1", "CONCRETOS", "Ensayos sobre muestras de concreto (cilindros, viguetas, cubos)"],
        ["2", "SUELOS", "Ensayos de clasificación y resistencia de suelos"],
        ["3", "ASFALTOS", "Ensayos sobre mezclas asfálticas"],
        ["4", "AGREGADOS", "Ensayos sobre materiales granulares"],
        ["5", "CEMENTOS", "Ensayos sobre muestras de cemento"],
        ["6", "OTROS", "Servicios especiales y consultoría"],
        ["7", "TRANSPORTE", "Servicio de transporte de muestras (regla especial de IVA)"],
    ]
)

doc.add_paragraph()

doc.add_heading('Crear una nueva categoría:', level=3)
step(1, "Clic en 'Nueva Categoría'", "Botón en la esquina superior derecha del catálogo.")
step(2, "Ingresar código", "Un número o código corto que identifique la categoría (ej: 1, 2, 3...).")
step(3, "Ingresar nombre", "Nombre descriptivo de la categoría (ej: CONCRETOS, SUELOS).")
step(4, "Guardar", "La categoría aparecerá en el catálogo inmediatamente.")

doc.add_paragraph()
doc.add_heading('Editar una categoría:', level=3)
doc.add_paragraph(
    'Use el botón "Editar" junto al nombre de la categoría en el catálogo. Puede modificar '
    'tanto el código como el nombre.'
)

doc.add_paragraph()
warning_box("La categoría con código '7' se considera automáticamente como TRANSPORTE. El sistema usa este código para aplicar la regla especial de IVA. No cambie el código de esta categoría.")

doc.add_heading('4.2 Tipos de Servicio', level=2)
doc.add_paragraph(
    'Cada tipo de servicio es un ensayo o servicio específico dentro de una categoría. '
    'Tiene un código único, nombre descriptivo y opcionalmente una norma técnica asociada.'
)

doc.add_heading('Ejemplo de servicios:', level=3)
table(
    ["Código", "Nombre del Servicio", "Categoría", "Norma"],
    [
        ["I-1", "Resistencia a la compresión de cilindros de concreto", "CONCRETOS", "NTC 673"],
        ["I-2", "Resistencia a la flexión de viguetas", "CONCRETOS", "NTC 2871"],
        ["II-1", "Granulometría por tamizado", "SUELOS", "INV E-123"],
        ["II-5", "Límites de Atterberg", "SUELOS", "INV E-125/126"],
        ["VII-1", "Transporte de muestras", "TRANSPORTE", "—"],
    ]
)

doc.add_paragraph()
doc.add_heading('Crear un nuevo servicio:', level=3)
step(1, "Clic en 'Nuevo Servicio'", "Botón en la esquina superior derecha del catálogo.")
step(2, "Seleccionar categoría", "Elija a qué categoría pertenece el servicio.")
step(3, "Ingresar código", "Código único del servicio (ej: I-1, II-3, VII-1).")
step(4, "Ingresar nombre", "Nombre completo y descriptivo del ensayo o servicio.")
step(5, "Ingresar norma (opcional)", "Norma técnica asociada (ej: NTC 673, INV E-123).")
step(6, "Guardar", "El servicio aparecerá dentro de su categoría en el catálogo.")

doc.add_paragraph()
doc.add_heading('Editar un servicio:', level=3)
doc.add_paragraph('Use el ícono de lápiz junto al servicio en el catálogo.')

doc.add_heading('Eliminar un servicio:', level=3)
doc.add_paragraph(
    'Use el ícono de papelera. El sistema pedirá confirmación. Solo se puede eliminar un '
    'servicio si no tiene registros de facturación asociados.'
)
warning_box("Si el servicio ya fue usado en algún registro, no se podrá eliminar para proteger la integridad del histórico.")

doc.add_heading('4.3 Buscar en el Catálogo', level=2)
doc.add_paragraph(
    'En la parte superior del catálogo hay un buscador que filtra por:'
)
bullet("Código del servicio (ej: 'I-1')")
bullet("Nombre del servicio (ej: 'compresión')")
bullet("Norma técnica (ej: 'NTC 673')")

doc.add_paragraph(
    'Escriba cualquier término y haga clic en buscar. Para volver a la vista completa, '
    'haga clic en "Limpiar".'
)

page_break()


# ============================================================
# 5. LISTA DE PRECIOS POR OBRA
# ============================================================
doc.add_heading('5. Lista de Precios por Obra', level=1)
doc.add_paragraph()
colored_p("Expediente de Obra → Botón $ (Lista de Precios)", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_heading('5.1 ¿Qué es la lista de precios?', level=2)
doc.add_paragraph(
    'Cada obra maneja sus propios precios para cada servicio. Esto permite que:'
)
bullet("Una obra pague $50,000 por un ensayo de compresión y otra pague $45,000.")
bullet("Se puedan negociar precios especiales por volumen o por contrato.")
bullet("Los precios de una obra sean completamente independientes de las demás.")

doc.add_paragraph()
doc.add_paragraph(
    'La lista de precios es el paso previo obligatorio antes de registrar servicios. '
    'Sin precios configurados, no se pueden registrar servicios para esa obra.'
)

doc.add_heading('5.2 Configurar precios paso a paso', level=2)

step(1, "Abrir el expediente de la obra", "Navegue a Directorio Empresas → seleccione la empresa → seleccione la obra.")
step(2, "Hacer clic en el botón '$'", "En el encabezado del expediente, junto a los botones de acción, verá el ícono de dólar ($). Haga clic.")
step(3, "La pantalla de precios se abre", "Verá todos los servicios del catálogo organizados por categoría, cada uno con un campo de precio.")
step(4, "Ingresar los precios", "Para cada servicio que aplique a esta obra, escriba el precio unitario en pesos. Deje en blanco los servicios que no aplican.")
step(5, "Guardar", "Haga clic en 'Guardar Precios'. Solo se guardarán los servicios con precio ingresado.")

doc.add_paragraph()

doc.add_heading('Ejemplo visual:', level=3)
table(
    ["Categoría", "Servicio", "Precio Unitario"],
    [
        ["CONCRETOS", "I-1 Resistencia a la compresión", "$25,000"],
        ["CONCRETOS", "I-2 Resistencia a la flexión", "$35,000"],
        ["SUELOS", "II-1 Granulometría por tamizado", "$40,000"],
        ["SUELOS", "II-5 Límites de Atterberg", "$30,000"],
        ["TRANSPORTE", "VII-1 Transporte de muestras", "$80,000"],
    ]
)

doc.add_paragraph()
tip_box("Servicios no aplicables", "Si un servicio del catálogo no se realizará en esta obra, simplemente deje el precio vacío. No aparecerá al registrar servicios.")

doc.add_heading('5.3 Precios y congelamiento', level=2)
doc.add_paragraph(
    'Un concepto fundamental del sistema es el congelamiento de precios:'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("¿Qué es el congelamiento?  ")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph(
    'Cuando usted registra un servicio para una obra, el sistema toma el precio vigente '
    'en ese momento y lo "congela" en el registro. Esto significa que:'
)
bullet("Si después cambia el precio en la lista de precios, los registros ya existentes NO se afectan.")
bullet("El precio congelado es el que aparecerá en la factura.")
bullet("Esto protege contra cambios de precio retroactivos y garantiza que cada factura refleje el precio acordado al momento de prestar el servicio.")

doc.add_paragraph()
doc.add_heading('Ejemplo práctico:', level=4)
doc.add_paragraph(
    '1. El 1 de marzo, el ensayo I-1 cuesta $25,000 para la obra "Torre Norte".\n'
    '2. Usted registra 5 ensayos I-1 el 1 de marzo → cada uno queda congelado a $25,000.\n'
    '3. El 15 de marzo, negocia un nuevo precio de $28,000 y lo actualiza en la lista.\n'
    '4. Registra 3 ensayos I-1 el 16 de marzo → estos quedan congelados a $28,000.\n'
    '5. Al facturar, los primeros 5 aparecen a $25,000 y los otros 3 a $28,000.\n\n'
    'Cada registro conserva su precio del momento exacto en que fue creado.'
)

page_break()


# ============================================================
# 6. GESTIÓN DE IMPUESTOS
# ============================================================
doc.add_heading('6. Gestión de Impuestos', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Impuestos", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Los impuestos se configuran de forma independiente y se aplican automáticamente al '
    'generar facturas. La pantalla de impuestos permite crear, editar y activar/desactivar '
    'los impuestos del sistema.'
)

doc.add_heading('6.1 Crear el IVA', level=2)
doc.add_paragraph(
    'Al instalar el sistema por primera vez, debe crear el impuesto IVA:'
)

step(1, "Ir a Impuestos", "Menú lateral → Impuestos.")
step(2, "Clic en 'Nuevo Impuesto'", "Se abre el formulario de creación.")
step(3, "Llenar los campos", "")
doc.add_paragraph()
table(
    ["Campo", "Valor", "Descripción"],
    [
        ["Nombre", "IVA", "Nombre del impuesto"],
        ["Porcentaje", "19.00", "Porcentaje actual del IVA en Colombia"],
        ["Activo", "✅ Marcado", "Indica que este impuesto está vigente"],
    ]
)
doc.add_paragraph()
step(4, "Guardar", "El IVA estará disponible para la generación de facturas.")

doc.add_paragraph()
tip_box("Cambio de tarifa", "Si la tarifa del IVA cambia en el futuro, simplemente edite el impuesto y modifique el porcentaje. Las facturas ya generadas no se verán afectadas.")

doc.add_heading('6.2 Regla Especial de Transporte', level=2)
doc.add_paragraph(
    'El sistema maneja una regla especial para los servicios de transporte:'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Regla de IVA:")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()
bullet("Los servicios normales (Concretos, Suelos, Asfaltos, etc.) SIEMPRE llevan IVA.")
bullet("Los servicios de TRANSPORTE (categoría código 7) solo llevan IVA si se selecciona explícitamente al generar la factura.")

doc.add_paragraph()
doc.add_paragraph(
    '¿Por qué? Porque el transporte de muestras puede o no estar gravado dependiendo de '
    'las condiciones del contrato con cada cliente. El sistema le da la flexibilidad de '
    'decidirlo factura por factura.'
)

doc.add_paragraph()
doc.add_heading('¿Cómo funciona en la práctica?', level=3)
doc.add_paragraph(
    'Al generar una factura que incluya servicios de transporte, estos aparecerán en una '
    'sección separada con un checkbox junto a cada uno. Marque el checkbox de los servicios '
    'de transporte que deben llevar IVA y deje sin marcar los que no. El total se recalcula '
    'en tiempo real.'
)

page_break()


# ============================================================
# 7. REGISTRO DE SERVICIOS
# ============================================================
doc.add_heading('7. Registro de Servicios Realizados', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Registrar Servicio", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El registro de servicios es la operación más frecuente del módulo. Cada vez que Geolab '
    'realiza un ensayo, prueba o servicio para una obra, se debe registrar aquí. Estos '
    'registros son la materia prima para generar las facturas.'
)

doc.add_heading('7.1 Proceso paso a paso', level=2)
doc.add_paragraph()

step(1, "Seleccionar la Constructora",
     "En el campo 'Constructora', empiece a escribir el nombre de la empresa. El buscador "
     "filtra las opciones automáticamente. Seleccione la empresa correcta.")

step(2, "Seleccionar la Obra",
     "Una vez seleccionada la constructora, el campo 'Obra' se habilita y carga las obras "
     "de esa empresa. Busque y seleccione la obra donde se prestó el servicio.")

step(3, "Indicar la Fecha de Realización",
     "Seleccione la fecha en que se realizó el servicio. Por defecto aparece la fecha de hoy, "
     "pero puede cambiarla si está registrando servicios de días anteriores.")

step(4, "Agregar líneas de servicio",
     "En la tabla de servicios, agregue una línea por cada tipo de servicio realizado:")

doc.add_paragraph()
doc.add_heading('Para cada línea:', level=4)

items_linea = [
    ("Servicio", "Busque el servicio escribiendo el código o nombre. El buscador filtra automáticamente entre todos los servicios del catálogo."),
    ("Nº Informe", "Escriba el número de informe de referencia. Este es el identificador del informe técnico asociado al servicio."),
    ("Cantidad", "Indique cuántas veces se realizó ese servicio (por defecto: 1)."),
    ("Precio Unitario", "Se muestra automáticamente al seleccionar el servicio. Es el precio configurado en la lista de precios de la obra."),
    ("Subtotal", "Se calcula automáticamente: Precio Unitario × Cantidad."),
]
for titulo, desc in items_linea:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_paragraph()

step(5, "Agregar más líneas (opcional)",
     "Haga clic en '+ Agregar línea' para registrar más servicios en la misma operación. "
     "Puede agregar tantas líneas como necesite.")

step(6, "Guardar Registros",
     "Haga clic en el botón 'Guardar Registros'. El sistema valida todos los datos y guarda "
     "todos los registros de forma atómica (todos o ninguno).")

doc.add_paragraph()
tip_box("Eficiencia", "Puede registrar todos los servicios del día para una obra en una sola operación agregando múltiples líneas. No necesita guardar uno por uno.")

doc.add_heading('7.2 Agregar múltiples líneas', level=2)
doc.add_paragraph(
    'El formulario permite agregar cualquier cantidad de líneas de servicio. Cada línea es '
    'independiente y puede tener un servicio, informe y cantidad diferentes.'
)

doc.add_heading('Ejemplo:', level=3)
doc.add_paragraph(
    'Suponga que hoy realizó para la obra "Torre Norte":\n'
    '- 3 ensayos de compresión de cilindros (informe GT-2026-001)\n'
    '- 1 ensayo de flexión de vigueta (informe GT-2026-002)\n'
    '- 1 servicio de transporte'
)

table(
    ["Servicio", "Nº Informe", "Cant.", "Precio Unit.", "Subtotal"],
    [
        ["I-1 Compresión cilindros", "GT-2026-001", "3", "$25,000", "$75,000"],
        ["I-2 Flexión viguetas", "GT-2026-002", "1", "$35,000", "$35,000"],
        ["VII-1 Transporte", "GT-2026-001", "1", "$80,000", "$80,000"],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Al guardar, se crean 3 registros independientes vinculados a la obra y la fecha.')

doc.add_heading('7.3 Validaciones y errores comunes', level=2)

doc.add_paragraph()
table(
    ["Situación", "Mensaje", "Solución"],
    [
        ["No seleccionó obra", "'Seleccione obra y fecha'", "Seleccione constructora y luego obra"],
        ["Servicio sin completar", "'Complete todos los campos'", "Llene servicio, Nº informe y cantidad en cada línea"],
        ["Sin líneas de servicio", "'Agregue al menos una línea'", "Use '+ Agregar línea' para crear al menos una"],
        ["Servicio sin precio", "Aparece 'Sin precio'", "Configure el precio en la lista de precios de la obra"],
        ["Obra sin precios", "Todos aparecen 'Sin precio'", "Vaya al expediente de obra → botón $ → configure precios"],
    ]
)

doc.add_paragraph()
doc.add_heading('Después de guardar:', level=3)
doc.add_paragraph(
    'El sistema muestra un mensaje verde de confirmación con la cantidad de registros creados. '
    'Las líneas se limpian automáticamente para que pueda seguir registrando si lo necesita. '
    'Los registros quedan con estado "Pendiente" hasta que se incluyan en una factura.'
)

page_break()


# ============================================================
# 8. HISTÓRICO DE REGISTROS
# ============================================================
doc.add_heading('8. Histórico de Registros', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Histórico", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El histórico es una vista completa de todos los registros de servicios del sistema. '
    'Permite consultar, filtrar, editar y eliminar registros que aún no han sido facturados.'
)

doc.add_heading('8.1 Filtros de búsqueda', level=2)
doc.add_paragraph(
    'En la parte superior encontrará cinco filtros que puede combinar:'
)

filtros = [
    ("Constructora", "Busque por nombre de empresa. Al seleccionar una, el campo de Obra se actualiza con sus obras."),
    ("Obra", "Filtre por obra específica. Solo muestra las obras de la constructora seleccionada."),
    ("Servicio", "Busque por tipo de servicio. Escriba código o nombre para filtrar."),
    ("Fecha Desde", "Inicio del rango de fechas a consultar."),
    ("Fecha Hasta", "Fin del rango de fechas a consultar."),
]
for titulo, desc in filtros:
    bullet_bold(titulo, desc)

doc.add_paragraph()
tip_box("Combinación", "Puede usar los filtros en cualquier combinación. Por ejemplo: solo constructora, o constructora + rango de fechas, o solo tipo de servicio.")

doc.add_heading('Información de cada registro:', level=3)
table(
    ["Columna", "Descripción"],
    [
        ["Fecha", "Fecha en que se realizó el servicio"],
        ["Constructora", "Empresa a la que pertenece la obra"],
        ["Obra", "Proyecto donde se prestó el servicio"],
        ["Servicio", "Código y nombre del tipo de servicio"],
        ["Nº Informe", "Referencia al informe técnico"],
        ["Cant.", "Cantidad de servicios realizados"],
        ["P. Unit.", "Precio unitario congelado al momento del registro"],
        ["Subtotal", "Precio unitario × Cantidad"],
        ["Estado", "Pendiente (amarillo) o Facturado #ID (verde)"],
    ]
)

doc.add_heading('8.2 Editar un registro', level=2)
doc.add_paragraph(
    'Solo puede editar registros que no hayan sido facturados (estado "Pendiente"). '
    'Haga clic en el ícono de lápiz para abrir el formulario de edición.'
)

doc.add_paragraph(
    'Al editar y guardar, el sistema actualiza el precio congelado al precio vigente actual '
    'de la lista de precios de la obra. Esto permite corregir errores de precio.'
)

warning_box("Los registros que ya están en una factura (estado 'Facturado #ID') no se pueden editar. Para modificarlos, primero anule la factura correspondiente.")

doc.add_heading('8.3 Eliminar un registro', level=2)
doc.add_paragraph(
    'Solo puede eliminar registros no facturados. Haga clic en el ícono de papelera '
    'y confirme la eliminación. Esta acción es irreversible.'
)

doc.add_paragraph()
doc.add_heading('Paginación:', level=3)
doc.add_paragraph(
    'Los registros se muestran de 25 en 25. Use los botones de paginación en la parte '
    'inferior para navegar entre páginas.'
)

page_break()


# ============================================================
# 9. GENERACIÓN DE FACTURAS
# ============================================================
doc.add_heading('9. Generación de Facturas', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Generar Factura", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'La generación de facturas es el proceso central del módulo. A partir de los registros '
    'pendientes de una obra en un periodo determinado, el sistema genera una factura formal '
    'con cálculo automático de impuestos y PDF profesional.'
)

doc.add_heading('9.1 Seleccionar obra y periodo', level=2)
doc.add_paragraph(
    'El primer paso es indicar qué obra quiere facturar y en qué rango de fechas:'
)

step(1, "Seleccionar Constructora", "Busque la empresa en el campo de constructora.")
step(2, "Seleccionar Obra", "Seleccione la obra específica.")
step(3, "Definir Fecha Inicio", "Primer día del periodo a facturar (ej: 1 de marzo).")
step(4, "Definir Fecha Fin", "Último día del periodo (ej: 31 de marzo).")
step(5, "Clic en 'Buscar registros'", "El sistema busca todos los registros pendientes de esa obra en ese periodo.")

doc.add_paragraph()
tip_box("Periodo típico", "Lo más común es facturar por mes (del 1 al 30/31). Pero puede usar cualquier rango que necesite.")
warning_box("Si no hay registros pendientes en el periodo seleccionado, el sistema mostrará un mensaje indicándolo. Verifique las fechas o que los servicios no estén ya facturados.")

doc.add_heading('9.2 Vista previa de la factura', level=2)
doc.add_paragraph(
    'Si hay registros pendientes, el sistema muestra una vista previa completa con todos '
    'los detalles antes de generar la factura. Esto le permite verificar que todo esté '
    'correcto antes de confirmar.'
)

doc.add_paragraph(
    'La vista previa incluye:'
)
bullet("Información de la obra (constructora, nombre, código)")
bullet("Periodo de facturación")
bullet("Tabla detallada de servicios con fecha, código, nombre, informe, cantidad, precio y subtotal")
bullet("Desglose de IVA")
bullet("Total de la factura")

doc.add_heading('9.3 Servicios normales vs. transporte', level=2)
doc.add_paragraph(
    'La vista previa separa los registros en dos secciones:'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Servicios Normales")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = BRAND_BLUE
doc.add_paragraph(
    'Todos los servicios que NO son de transporte. Estos siempre llevan IVA. Se muestran '
    'en una tabla con subtotal al final.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Servicios de Transporte")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = WARNING_ORANGE
doc.add_paragraph(
    'Los servicios de la categoría 7 (Transporte) se muestran en una tabla separada con '
    'fondo amarillo. Cada servicio tiene un checkbox para indicar si lleva IVA o no.'
)

doc.add_heading('9.4 Selección de IVA en transporte', level=2)
doc.add_paragraph(
    'Para cada servicio de transporte, debe decidir si lleva IVA:'
)

bullet("Checkbox MARCADO: El transporte se incluye en la base para calcular IVA.")
bullet("Checkbox SIN MARCAR: El transporte NO lleva IVA (se factura sin impuesto).")

doc.add_paragraph(
    'Al marcar o desmarcar los checkboxes, el resumen de la factura (IVA y Total) se '
    'recalcula en tiempo real.'
)

doc.add_paragraph()
doc.add_heading('Ejemplo de cálculo:', level=3)
doc.add_paragraph(
    'Suponga que la factura tiene:'
)

table(
    ["Concepto", "Subtotal", "¿Lleva IVA?"],
    [
        ["Servicios normales", "$500,000", "Siempre (automático)"],
        ["Transporte viaje 1", "$80,000", "✅ Sí (marcado)"],
        ["Transporte viaje 2", "$80,000", "❌ No (sin marcar)"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Cálculo del IVA (19%):\n'
    '• Base gravable = $500,000 (servicios) + $80,000 (transporte con IVA) = $580,000\n'
    '• IVA = $580,000 × 19% = $110,200\n'
    '• Subtotal general = $500,000 + $80,000 + $80,000 = $660,000\n'
    '• TOTAL = $660,000 + $110,200 = $770,200'
)

doc.add_heading('9.5 Confirmar y generar', level=2)
doc.add_paragraph(
    'Una vez verificada la vista previa:'
)

step(1, "Revisar los datos", "Verifique que todos los servicios, cantidades y precios sean correctos.")
step(2, "Decidir IVA de transporte", "Marque los checkboxes correspondientes si hay servicios de transporte.")
step(3, "Clic en 'Generar Factura'", "El sistema pide confirmación.")
step(4, "Confirmar", "Acepte el diálogo de confirmación.")

doc.add_paragraph()
doc.add_paragraph('Al confirmar, el sistema automáticamente:')
bullet("Calcula el IVA total según las reglas (normales + transporte seleccionado)")
bullet("Crea el registro de la factura con todos los montos")
bullet("Vincula todos los registros de servicio a la factura (cambian de 'Pendiente' a 'Facturado')")
bullet("Genera un PDF profesional con el formato de Geolab")
bullet("Redirige al repositorio de facturas de esa obra")

doc.add_heading('9.6 PDF de la factura', level=2)
doc.add_paragraph(
    'El PDF generado automáticamente incluye:'
)
bullet("Encabezado con datos de Geolab")
bullet("Datos del cliente (constructora, NIT)")
bullet("Datos de la obra (nombre, código)")
bullet("Periodo de facturación")
bullet("Tabla detallada de servicios con cantidades y precios")
bullet("Desglose: Subtotal, IVA, Total")
bullet("Número de factura y fecha de emisión")

doc.add_paragraph()
tip_box("Formato", "El PDF se genera en tamaño carta con los colores corporativos de Geolab, listo para imprimir o enviar por email.")

page_break()


# ============================================================
# 10. REPOSITORIO DE FACTURAS
# ============================================================
doc.add_heading('10. Repositorio de Facturas', level=1)
doc.add_paragraph()
colored_p("Menú lateral → Repositorio", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El repositorio organiza todas las facturas generadas en una navegación jerárquica '
    'de tres niveles, desde los clientes hasta las facturas individuales.'
)

doc.add_heading('10.1 Nivel 1: Clientes', level=2)
doc.add_paragraph(
    'La pantalla inicial del repositorio muestra tarjetas con todas las empresas constructoras '
    'que tienen facturas. Cada tarjeta muestra:'
)
bullet("Nombre de la empresa")
bullet("Cantidad de facturas generadas")
bullet("Monto total facturado")

doc.add_paragraph('Haga clic en una empresa para ver sus obras.')

doc.add_heading('10.2 Nivel 2: Obras del Cliente', level=2)
doc.add_paragraph(
    'Al seleccionar un cliente, ve la lista de sus obras que tienen facturas. '
    'Cada obra muestra el conteo de facturas y el monto facturado.'
)
doc.add_paragraph('Haga clic en una obra para ver sus facturas.')

doc.add_heading('10.3 Nivel 3: Facturas de la Obra', level=2)
doc.add_paragraph(
    'El nivel final muestra la tabla detallada de todas las facturas de una obra:'
)

table(
    ["Columna", "Descripción"],
    [
        ["#", "Número de identificación de la factura"],
        ["Fecha emisión", "Fecha en que se generó la factura"],
        ["Periodo", "Rango de fechas del periodo facturado"],
        ["Subtotal", "Suma de los servicios sin impuestos"],
        ["IVA", "Monto del IVA calculado"],
        ["Total", "Monto total de la factura (subtotal + IVA)"],
        ["Estado", "Pendiente (amarillo), Pagada (verde), Anulada (rojo)"],
        ["Acciones", "Botones: Ver PDF, Marcar pagada, Anular"],
    ]
)

doc.add_paragraph()
tip_box("Facturas anuladas", "Las facturas anuladas se muestran tachadas y en color rojo, pero siguen visibles como registro histórico.")

page_break()


# ============================================================
# 11. ACCIONES SOBRE FACTURAS
# ============================================================
doc.add_heading('11. Acciones sobre Facturas', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'Desde el repositorio (nivel 3) puede realizar acciones sobre cada factura según su estado.'
)

doc.add_heading('11.1 Marcar como Pagada', level=2)
doc.add_paragraph(
    'Cuando el cliente paga una factura, use el botón verde (✓) para cambiar su estado:'
)
bullet("Solo funciona con facturas en estado PENDIENTE.")
bullet("Al hacer clic, el sistema pide confirmación: '¿Marcar como pagada?'")
bullet("El estado cambia a PAGADA (verde) y los botones de acción se ocultan.")
bullet("Esta acción no se puede deshacer directamente.")

doc.add_heading('11.2 Anular una Factura', level=2)
doc.add_paragraph(
    'Si una factura fue generada por error o necesita ser corregida, puede anularla:'
)

step(1, "Clic en el botón rojo (✗)", "Solo disponible para facturas PENDIENTES.")
step(2, "Confirmar", "El sistema pregunta: '¿Anular esta factura? Los registros quedarán disponibles para re-facturación.'")
step(3, "Aceptar", "")

doc.add_paragraph()
doc.add_paragraph('Al anular, el sistema:')
bullet("Marca la factura como ANULADA (no la elimina)")
bullet("Desvincula todos los registros de servicio de esa factura")
bullet("Los registros vuelven a estado 'Pendiente' y quedan disponibles para incluirlos en una nueva factura")
bullet("La factura anulada permanece visible en el repositorio como registro histórico, mostrada tachada y en rojo")

doc.add_paragraph()
doc.add_heading('¿Cuándo anular?', level=3)
doc.add_paragraph('Casos típicos de anulación:')
bullet("Se incluyeron servicios que no correspondían")
bullet("Se seleccionó mal el periodo de facturación")
bullet("Error en la selección de IVA en servicios de transporte")
bullet("El cliente solicita re-facturación con cambios")

doc.add_paragraph()
warning_box("Solo se pueden anular facturas en estado PENDIENTE. Las facturas PAGADAS no se pueden anular. Si necesita corregir una factura pagada, contacte al administrador del sistema.")

doc.add_heading('11.3 Descargar PDF', level=2)
doc.add_paragraph(
    'Use el botón azul de PDF para descargar o visualizar el documento de la factura. '
    'El PDF se abre en una nueva pestaña del navegador desde donde puede imprimirlo o guardarlo.'
)

doc.add_heading('11.4 Resumen de acciones por estado', level=2)
doc.add_paragraph()

table(
    ["Estado", "Ver PDF", "Marcar Pagada", "Anular"],
    [
        ["PENDIENTE", "✅", "✅", "✅"],
        ["PAGADA", "✅", "— (ya pagada)", "❌"],
        ["ANULADA", "✅ (histórico)", "❌", "— (ya anulada)"],
    ]
)

page_break()


# ============================================================
# 12. EJEMPLO PRÁCTICO
# ============================================================
doc.add_heading('12. Flujo Completo: Ejemplo Práctico', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'A continuación se muestra un ejemplo completo del flujo de facturación para la obra '
    '"Edificio Torres del Parque" de la empresa "Constructora Bolívar S.A."'
)

doc.add_heading('Escenario', level=3)
doc.add_paragraph(
    'Es marzo de 2026. Geolab realizó diversos ensayos para la obra durante el mes y necesita '
    'generar la factura mensual.'
)

doc.add_paragraph()
doc.add_heading('Paso 1: Verificar el catálogo (una sola vez)', level=2)
doc.add_paragraph(
    'Antes de comenzar, verifique que los servicios que va a registrar existan en el catálogo. '
    'En este caso ya existen: I-1 (Compresión), I-2 (Flexión), VII-1 (Transporte).'
)

doc.add_heading('Paso 2: Configurar precios de la obra (una sola vez)', level=2)
doc.add_paragraph(
    'Vaya al expediente de la obra → botón $ → Configure:'
)
table(
    ["Servicio", "Precio"],
    [
        ["I-1 Compresión cilindros", "$25,000"],
        ["I-2 Flexión viguetas", "$35,000"],
        ["VII-1 Transporte", "$80,000"],
    ]
)

doc.add_heading('Paso 3: Registrar servicios del mes', level=2)
doc.add_paragraph('Durante marzo, registra los servicios a medida que se realizan:')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("5 de marzo:")
run.bold = True
doc.add_paragraph('→ 4 ensayos de compresión (informe GT-2026-015) + 1 transporte')

p = doc.add_paragraph()
run = p.add_run("12 de marzo:")
run.bold = True
doc.add_paragraph('→ 6 ensayos de compresión (informe GT-2026-018) + 2 flexiones (informe GT-2026-019)')

p = doc.add_paragraph()
run = p.add_run("20 de marzo:")
run.bold = True
doc.add_paragraph('→ 3 ensayos de compresión (informe GT-2026-022) + 1 transporte')

doc.add_heading('Paso 4: Generar la factura', level=2)
doc.add_paragraph('A fin de mes, vaya a Generar Factura:')
bullet("Constructora: Constructora Bolívar S.A.")
bullet("Obra: Edificio Torres del Parque")
bullet("Fecha inicio: 01/03/2026")
bullet("Fecha fin: 31/03/2026")
bullet("Clic en Buscar registros")

doc.add_paragraph()
doc.add_paragraph('La vista previa muestra:')
doc.add_paragraph()

table(
    ["Sección", "Detalle", "Subtotal"],
    [
        ["Servicios normales", "13 compresiones × $25,000 + 2 flexiones × $35,000", "$395,000"],
        ["Transporte", "2 transportes × $80,000", "$160,000"],
        ["", "SUBTOTAL GENERAL", "$555,000"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Decide que los dos transportes llevan IVA → marca ambos checkboxes.\n'
    '• Base IVA = $395,000 + $160,000 = $555,000\n'
    '• IVA (19%) = $105,450\n'
    '• TOTAL = $660,450'
)

doc.add_heading('Paso 5: Confirmar', level=2)
doc.add_paragraph(
    'Clic en "Generar Factura" → Confirmar → El sistema genera la factura, el PDF y '
    'redirige al repositorio donde puede descargar el documento.'
)

doc.add_heading('Paso 6: Gestionar', level=2)
doc.add_paragraph('Cuando el cliente pague:')
bullet("Vaya al Repositorio → Constructora Bolívar → Torres del Parque")
bullet("En la factura, haga clic en el botón verde (✓) → Confirmar")
bullet("La factura pasa a estado PAGADA")

page_break()


# ============================================================
# 13. PREGUNTAS FRECUENTES
# ============================================================
doc.add_heading('13. Preguntas Frecuentes', level=1)
doc.add_paragraph()

faqs = [
    ("¿Puedo facturar servicios de varias obras en una sola factura?",
     "No. Cada factura corresponde a una sola obra y un periodo específico. Si necesita "
     "facturar varias obras del mismo cliente, genere una factura por cada obra."),

    ("¿Qué pasa si me equivoqué en un registro ya facturado?",
     "Debe anular la factura primero. Al anularla, los registros quedan libres para "
     "editarlos o eliminarlos. Luego genere una nueva factura con los datos correctos."),

    ("¿Por qué un servicio aparece como 'Sin precio'?",
     "Porque no se ha configurado el precio de ese servicio para esa obra. Vaya al "
     "expediente de la obra → botón $ → ingrese el precio del servicio faltante."),

    ("¿Si cambio el precio de un servicio, se afectan los registros existentes?",
     "No. Los registros conservan el precio 'congelado' del momento en que fueron creados. "
     "El cambio de precio solo aplica a nuevos registros futuros."),

    ("¿Puedo generar una factura sin IVA?",
     "Los servicios normales siempre llevan IVA. Solo los servicios de transporte permiten "
     "elegir si llevan IVA o no. Si necesita una factura completamente sin IVA, solo "
     "los servicios de transporte lo permiten por diseño."),

    ("¿Puedo re-generar el PDF de una factura?",
     "El PDF se genera una sola vez al crear la factura y se almacena en el sistema. "
     "Puede descargarlo cuantas veces necesite desde el repositorio."),

    ("¿Qué significa 'registros pendientes'?",
     "Son servicios que han sido registrados pero aún no se han incluido en ninguna factura. "
     "Es decir, son servicios realizados y cobrados pero no formalizados en una factura."),

    ("¿Puedo eliminar una factura completamente?",
     "No. Las facturas solo se pueden anular, no eliminar. Las facturas anuladas permanecen "
     "como registro histórico para trazabilidad contable."),

    ("¿Quién puede acceder al módulo de facturación?",
     "Solo el Staff Administrativo (usuarios con rol admin, laboratorio o recepción en Geolab). "
     "Los técnicos de laboratorio y los clientes externos no tienen acceso."),

    ("¿Debo crear el IVA manualmente?",
     "Sí. Al usar el módulo por primera vez, vaya a Impuestos → Nuevo Impuesto → cree el IVA "
     "con porcentaje 19%. Solo necesita hacerlo una vez."),

    ("¿Puedo registrar servicios de días anteriores?",
     "Sí. El campo de fecha permite seleccionar cualquier fecha pasada. Es común registrar "
     "los servicios de la semana o del día anterior."),

    ("¿Cuántos servicios puedo agregar en un solo registro?",
     "No hay límite. Puede agregar tantas líneas como necesite usando el botón '+ Agregar línea'."),

    ("¿Puedo facturar un periodo parcial?",
     "Sí. Puede seleccionar cualquier rango de fechas. No tiene que ser un mes completo. "
     "Por ejemplo, puede facturar del 1 al 15 de marzo si así lo necesita."),

    ("¿Qué pasa con los registros si anulo la factura?",
     "Vuelven a estado 'Pendiente'. Puede editarlos, eliminarlos o incluirlos en una nueva "
     "factura. No se pierden."),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    run = p.add_run(pregunta)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(respuesta)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(10)

page_break()


# ============================================================
# 14. GLOSARIO
# ============================================================
doc.add_heading('14. Glosario de Términos', level=1)
doc.add_paragraph()

glosario = [
    ("Categoría de Servicio", "Agrupación temática de servicios (ej: Concretos, Suelos, Transporte)."),
    ("Tipo de Servicio", "Servicio o ensayo específico dentro de una categoría (ej: Compresión de cilindros)."),
    ("Lista de Precios", "Conjunto de precios asignados a cada servicio para una obra específica."),
    ("Precio Congelado", "Precio que se fija en el momento del registro y no cambia aunque se modifique la lista de precios."),
    ("Registro de Servicio", "Entrada individual que documenta un servicio realizado para una obra en una fecha."),
    ("Registro Pendiente", "Servicio registrado que aún no ha sido incluido en ninguna factura."),
    ("Factura", "Documento formal que agrupa registros de servicio de una obra en un periodo, con cálculo de impuestos."),
    ("Periodo de Facturación", "Rango de fechas que abarca una factura (típicamente un mes)."),
    ("IVA", "Impuesto al Valor Agregado, actualmente 19% en Colombia."),
    ("Regla de Transporte", "Norma del sistema donde los servicios de transporte (categoría 7) solo llevan IVA si se marca explícitamente."),
    ("Anulación", "Proceso de cancelar una factura, liberando los registros para re-facturación. La factura no se elimina."),
    ("Repositorio", "Archivo organizado de todas las facturas, navegable por cliente → obra → factura."),
    ("Constructora", "Empresa constructora cliente de Geolab."),
    ("Obra", "Proyecto de construcción perteneciente a una constructora."),
    ("Número de Informe", "Referencia al informe técnico de Geolab asociado al servicio registrado."),
    ("Subtotal", "Suma de los servicios sin impuestos."),
    ("Base Gravable", "Monto sobre el cual se calcula el IVA (servicios normales + transporte con IVA seleccionado)."),
]

for termino, definicion in glosario:
    p = doc.add_paragraph()
    run = p.add_run(f"{termino}: ")
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run2 = p.add_run(definicion)
    p.paragraph_format.space_after = Pt(6)


page_break()

# ============================================================
# CONTRAPORTADA
# ============================================================
for _ in range(8):
    doc.add_paragraph()

colored_p("GEOLAB S.A.S", BRAND_BLUE, bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("Módulo de Facturación - Manual de Usuario v1.0", DARK_BLUE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
colored_p("Marzo 2026", GRAY_TEXT, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

colored_p("Para soporte técnico contacte al equipo de Geolab S.A.S", GRAY_TEXT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
colored_p("https://portalgeolab.com", BRAND_BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()
colored_p("Desarrollado por Vadom Data Consulting SAS", GRAY_TEXT, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================
# GUARDAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Manual_Facturacion_Geolab.docx')
doc.save(output_path)
print(f"\n✅ Manual de Facturación generado: {output_path}")
print(f"   Tamaño: {os.path.getsize(output_path) / 1024:.0f} KB")
