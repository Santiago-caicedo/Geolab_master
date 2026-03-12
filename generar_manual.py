"""
Script para generar el Manual de Usuario de Geolab Portal
Ejecutar: python3 generar_manual.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

# ============================================================
# CONFIGURACIÓN
# ============================================================
BRAND_BLUE = RGBColor(0, 142, 214)       # #008ed6
DARK_BLUE = RGBColor(28, 46, 74)         # #1c2e4a
LIGHT_GRAY = RGBColor(248, 250, 252)     # #f8fafc
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
GRAY_TEXT = RGBColor(100, 116, 139)       # #64748b
SUCCESS_GREEN = RGBColor(22, 163, 74)
WARNING_ORANGE = RGBColor(234, 88, 12)
DANGER_RED = RGBColor(220, 38, 38)

doc = Document()

# ============================================================
# ESTILOS BASE
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(51, 65, 85)

# Estilos de encabezados
for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = DARK_BLUE
    if i == 1:
        h.font.size = Pt(26)
        h.font.bold = True
    elif i == 2:
        h.font.size = Pt(20)
        h.font.bold = True
    elif i == 3:
        h.font.size = Pt(15)
        h.font.bold = True
    elif i == 4:
        h.font.size = Pt(12)
        h.font.bold = True

# ============================================================
# FUNCIONES AUXILIARES
# ============================================================
def add_colored_paragraph(text, color=None, bold=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_tip_box(title, text, icon="💡"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {icon} {title}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_BLUE
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY_TEXT

def add_warning_box(text):
    add_tip_box("Importante", text, icon="⚠️")

def add_step(number, title, description):
    p = doc.add_paragraph()
    run_num = p.add_run(f"  Paso {number}: ")
    run_num.bold = True
    run_num.font.color.rgb = BRAND_BLUE
    run_num.font.size = Pt(11)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    if description:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        run_desc = p2.add_run(description)
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = GRAY_TEXT

def add_simple_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def add_page_break():
    doc.add_page_break()

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 70)
    run.font.color.rgb = RGBColor(226, 232, 240)
    run.font.size = Pt(8)


# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_colored_paragraph("GEOLAB S.A.S", BRAND_BLUE, bold=True, size=36,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
add_colored_paragraph("Portal de Gestión Técnica", DARK_BLUE, bold=True, size=22,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_colored_paragraph("MANUAL DE USUARIO", DARK_BLUE, bold=True, size=28,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

add_colored_paragraph("Guía completa para el uso del sistema", GRAY_TEXT, size=14,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_colored_paragraph("Versión 1.0  |  Marzo 2026", GRAY_TEXT, size=11,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
add_colored_paragraph("Desarrollado por Vadom Data Consulting SAS", GRAY_TEXT, size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# TABLA DE CONTENIDO
# ============================================================
doc.add_heading('Tabla de Contenido', level=1)
doc.add_paragraph()

toc_items = [
    ("1.", "Introducción al Sistema"),
    ("2.", "Acceso al Sistema (Inicio de Sesión)"),
    ("3.", "Roles y Perfiles de Usuario"),
    ("4.", "Panel Principal (Dashboard)"),
    ("5.", "Directorio de Empresas y Obras"),
    ("6.", "Gestión de Informes Técnicos"),
    ("7.", "Remisiones de Muestras (F-GC-05)"),
    ("8.", "Ensayos de Laboratorio (F-GT-05)"),
    ("9.", "Sistema de Facturación"),
    ("10.", "Sistema de Calidad"),
    ("11.", "Gestión de Usuarios y Accesos"),
    ("12.", "Notificaciones"),
    ("13.", "Preguntas Frecuentes"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run_n = p.add_run(f"{num}  ")
    run_n.bold = True
    run_n.font.color.rgb = BRAND_BLUE
    run_n.font.size = Pt(12)
    run_t = p.add_run(title)
    run_t.font.size = Pt(12)

add_page_break()


# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
doc.add_heading('1. Introducción al Sistema', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'El Portal Geolab es la plataforma digital de Geolab S.A.S, laboratorio especializado en '
    'ensayos de suelos y materiales de construcción. Este sistema centraliza toda la operación '
    'del laboratorio: desde la recepción de muestras hasta la entrega de informes técnicos, '
    'pasando por el control de ensayos, la facturación y el sistema de gestión de calidad.'
)

doc.add_heading('¿Qué puede hacer con este sistema?', level=3)

beneficios = [
    ("Gestión de Clientes y Proyectos", "Organice toda la información de sus empresas constructoras y sus obras en un solo lugar."),
    ("Remisiones Digitales", "Reciba y gestione las remisiones de muestras (formato F-GC-05) de forma digital, eliminando el papel."),
    ("Control de Ensayos", "Los técnicos de laboratorio registran resultados en tiempo real, con cálculos automáticos."),
    ("Informes Técnicos", "Cargue, organice y distribuya los informes PDF a cada obra de forma segura."),
    ("Facturación", "Registre servicios prestados, genere facturas automáticas y lleve el histórico completo."),
    ("Calidad", "Gestione toda la documentación del Sistema de Gestión de Calidad organizada por áreas."),
    ("Acceso por Roles", "Cada usuario ve exactamente lo que necesita según su rol en la organización."),
]

for titulo, desc in beneficios:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)

doc.add_paragraph()

doc.add_heading('Acceso al Portal', level=3)
doc.add_paragraph(
    'El portal está disponible en la dirección https://portalgeolab.com desde cualquier '
    'navegador web (Google Chrome, Firefox, Edge, Safari). Es compatible con computadores, '
    'tablets y celulares.'
)

add_tip_box("Recomendación", "Use Google Chrome o Microsoft Edge para la mejor experiencia.")

add_page_break()


# ============================================================
# 2. ACCESO AL SISTEMA
# ============================================================
doc.add_heading('2. Acceso al Sistema', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'Para ingresar al sistema necesita un usuario y contraseña proporcionados por el '
    'administrador de Geolab. No existe opción de auto-registro; todos los usuarios son '
    'creados por el equipo administrativo.'
)

doc.add_heading('Inicio de Sesión', level=2)

add_step(1, "Abrir el navegador", "Ingrese a https://portalgeolab.com en la barra de direcciones.")
add_step(2, "Ingresar credenciales", "Escriba su nombre de usuario y contraseña en los campos correspondientes.")
add_step(3, "Hacer clic en \"Iniciar Sesión\"", "El sistema lo redirigirá automáticamente a su panel según su rol.")

doc.add_paragraph()
add_warning_box(
    "Si olvidó su contraseña, contacte al administrador de Geolab para que la restablezca. "
    "No comparta sus credenciales con otras personas."
)

doc.add_heading('¿A dónde me lleva el sistema al iniciar sesión?', level=3)

add_simple_table(
    ["Rol", "Destino al iniciar sesión"],
    [
        ["Staff Administrativo", "Panel Principal con estadísticas generales y actividad reciente"],
        ["Técnico de Laboratorio", "Panel de Ensayos con muestras programadas para hoy"],
        ["Director (Cliente)", "Portal de Cliente con las obras de su empresa"],
        ["Residente (Cliente)", "Portal de Cliente con sus obras asignadas"],
        ["Remitente (Cliente)", "Panel de Remitente para crear y consultar remisiones"],
    ]
)

add_page_break()


# ============================================================
# 3. ROLES Y PERFILES
# ============================================================
doc.add_heading('3. Roles y Perfiles de Usuario', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'El sistema maneja cinco tipos de usuario, cada uno con accesos y funcionalidades específicas. '
    'Esto garantiza que cada persona vea solo la información relevante para su trabajo.'
)

doc.add_heading('3.1 Staff Administrativo (Geolab)', level=2)
doc.add_paragraph(
    'Personal interno de Geolab con acceso completo al sistema. Incluye roles de administración, '
    'laboratorio y recepción.'
)
p = doc.add_paragraph()
run = p.add_run("Funciones principales:")
run.bold = True
items_admin = [
    "Gestionar el directorio de empresas y obras",
    "Cargar informes técnicos (PDFs)",
    "Crear y administrar remisiones de muestras",
    "Acceder a hojas de trabajo y ensayos",
    "Registrar servicios y generar facturas",
    "Gestionar documentación de calidad",
    "Crear y administrar usuarios del sistema",
    "Ver el panel de estadísticas generales",
]
for item in items_admin:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Técnico de Laboratorio', level=2)
doc.add_paragraph(
    'Personal de laboratorio de Geolab. Su función principal es registrar los resultados de los '
    'ensayos sobre las muestras recibidas.'
)
p = doc.add_paragraph()
run = p.add_run("Funciones principales:")
run.bold = True
items_tec = [
    "Ver las muestras programadas para ensayo del día",
    "Registrar mediciones (diámetros, longitudes, peso, carga, forma de falla)",
    "El sistema calcula automáticamente: esfuerzo, porcentaje de desarrollo y cumplimiento",
    "Marcar muestras como completadas",
    "Consultar documentación del sistema de calidad",
]
for item in items_tec:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 Director (Cliente Externo)', level=2)
doc.add_paragraph(
    'Representante de una empresa constructora con visibilidad total sobre todas las obras '
    'de su empresa. Es el perfil con mayor acceso entre los clientes.'
)
p = doc.add_paragraph()
run = p.add_run("Funciones principales:")
run.bold = True
items_dir = [
    "Ver todas las obras de su empresa",
    "Descargar los informes técnicos de cualquier obra",
    "Consultar las remisiones de muestras y su estado",
    "Buscar y filtrar informes por fecha o título",
]
for item in items_dir:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.4 Residente (Cliente Externo)', level=2)
doc.add_paragraph(
    'Ingeniero residente de obra. Solo tiene acceso a las obras específicas donde ha sido '
    'asignado por el administrador, no a todas las obras de la empresa.'
)
p = doc.add_paragraph()
run = p.add_run("Funciones principales:")
run.bold = True
items_res = [
    "Ver únicamente las obras donde está asignado",
    "Descargar informes técnicos de sus obras",
    "Consultar remisiones de sus obras",
]
for item in items_res:
    doc.add_paragraph(item, style='List Bullet')

add_tip_box("Diferencia clave", "El Director ve TODAS las obras de la empresa. El Residente solo ve las que le asignan.")

doc.add_heading('3.5 Remitente (Cliente Externo)', level=2)
doc.add_paragraph(
    'Persona encargada de enviar las muestras de campo al laboratorio. Su función principal '
    'es crear remisiones de muestras digitales (formato F-GC-05) desde el campo.'
)
p = doc.add_paragraph()
run = p.add_run("Funciones principales:")
run.bold = True
items_rem = [
    "Crear remisiones de muestras para las obras asignadas",
    "Registrar los detalles de cada muestra (tipo, localización, fecha de toma, etc.)",
    "Consultar el historial de sus remisiones enviadas",
    "Ver el estado de cada remisión (completada, en proceso)",
]
for item in items_rem:
    doc.add_paragraph(item, style='List Bullet')

add_page_break()

# ============================================================
# 3.6 Resumen de accesos
# ============================================================
doc.add_heading('3.6 Resumen Comparativo de Accesos', level=2)
doc.add_paragraph()

add_simple_table(
    ["Funcionalidad", "Staff Admin", "Técnico", "Director", "Residente", "Remitente"],
    [
        ["Panel Principal / Dashboard", "✅", "✅", "✅", "✅", "✅"],
        ["Directorio de Empresas", "✅", "—", "—", "—", "—"],
        ["Gestión de Obras", "✅", "—", "Solo lectura", "Solo lectura", "—"],
        ["Cargar Informes", "✅", "—", "—", "—", "—"],
        ["Descargar Informes", "✅", "—", "✅ (empresa)", "✅ (asignadas)", "—"],
        ["Crear Remisiones", "✅", "—", "—", "—", "✅"],
        ["Ver Remisiones", "✅", "—", "✅", "✅", "✅ (propias)"],
        ["Ensayos de Lab", "✅", "✅", "—", "—", "—"],
        ["Hojas de Trabajo", "✅", "✅", "—", "—", "—"],
        ["Facturación", "✅", "—", "—", "—", "—"],
        ["Sistema de Calidad", "✅", "✅ (asignado)", "—", "—", "—"],
        ["Gestión de Usuarios", "✅", "—", "—", "—", "—"],
        ["Notificaciones", "✅", "✅", "✅", "✅", "✅"],
    ]
)

add_page_break()


# ============================================================
# 4. PANEL PRINCIPAL
# ============================================================
doc.add_heading('4. Panel Principal (Dashboard)', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'El Panel Principal es la primera pantalla que ve al iniciar sesión. Su contenido varía '
    'según el rol del usuario.'
)

doc.add_heading('4.1 Panel del Staff Administrativo', level=2)
doc.add_paragraph(
    'El panel del staff muestra una vista general de toda la operación del laboratorio en '
    'tiempo real.'
)

doc.add_heading('Indicadores Principales (KPIs)', level=3)
kpis_staff = [
    ("Total de Empresas", "Cantidad de empresas constructoras registradas en el sistema."),
    ("Total de Obras", "Número de proyectos/obras activos."),
    ("Informes Cargados", "Cantidad total de informes técnicos en el sistema."),
    ("Remisiones Recibidas", "Total de remisiones de muestras registradas."),
]
for titulo, desc in kpis_staff:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run2 = p.add_run(desc)

doc.add_heading('Secciones del Panel', level=3)
secciones_panel = [
    ("Muestras Programadas para Hoy",
     "Muestra las muestras que deben ser ensayadas hoy según la fecha de toma + días de edad. "
     "Las muestras vencidas aparecen resaltadas en rojo como alerta."),
    ("Últimos Informes Cargados",
     "Los 8 informes más recientes cargados al sistema, con la obra y fecha."),
    ("Últimas Remisiones Recibidas",
     "Las 5 remisiones más recientes con su estado (completada, pendiente, vencida)."),
    ("Obras con Mayor Actividad",
     "Las 5 obras con más movimiento en los últimos 30 días."),
]
for titulo, desc in secciones_panel:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(0.5)

doc.add_heading('4.2 Panel del Cliente (Director/Residente)', level=2)
doc.add_paragraph(
    'Los clientes ven un panel simplificado centrado en sus obras y los informes disponibles.'
)

items_client = [
    "Tarjetas de sus obras con conteo de informes disponibles",
    "Acceso rápido a los últimos 5 informes cargados",
    "Total de informes disponibles para consulta",
    "Acceso directo a cada obra para ver sus informes",
]
for item in items_client:
    doc.add_paragraph(item, style='List Bullet')

add_tip_box("Nota", "El Director ve todas las obras de su empresa. El Residente solo ve las obras donde está asignado.")

doc.add_heading('4.3 Panel del Remitente', level=2)
doc.add_paragraph(
    'El remitente ve un panel enfocado en la creación y seguimiento de remisiones de muestras.'
)
items_remitente = [
    "Lista de obras asignadas con botón para crear nueva remisión",
    "Estadísticas: total de remisiones enviadas, pendientes y completadas",
    "Últimas 5 remisiones creadas con enlace al detalle",
]
for item in items_remitente:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 Panel del Técnico de Laboratorio', level=2)
doc.add_paragraph(
    'El técnico ve un panel orientado a las tareas diarias del laboratorio.'
)
items_tec_panel = [
    "Muestras programadas para ensayar hoy",
    "Muestras vencidas pendientes de ensayo",
    "Resumen diario: pendientes y completados",
    "Acceso directo a las hojas de trabajo",
]
for item in items_tec_panel:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.5 Barra de Navegación Lateral (Sidebar)', level=2)
doc.add_paragraph(
    'A la izquierda de la pantalla encontrará el menú de navegación lateral. Este menú muestra '
    'diferentes opciones según su rol. Puede ocultarlo o mostrarlo haciendo clic en el ícono '
    'de hamburguesa (tres líneas) en la esquina superior izquierda.'
)

doc.add_heading('4.6 Barra Superior', level=2)
doc.add_paragraph(
    'En la parte superior encontrará:'
)
items_barra = [
    "Botón para ocultar/mostrar el menú lateral",
    "Nombre de su panel según su rol",
    "Fecha actual",
    "Campana de notificaciones con contador de no leídas",
    "Su nombre de usuario e inicial",
]
for item in items_barra:
    doc.add_paragraph(item, style='List Bullet')

add_page_break()


# ============================================================
# 5. DIRECTORIO DE EMPRESAS Y OBRAS
# ============================================================
doc.add_heading('5. Directorio de Empresas y Obras', level=1)
doc.add_paragraph()
add_colored_paragraph("Acceso: Solo Staff Administrativo", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El directorio es el punto central para administrar las empresas constructoras (clientes) '
    'y sus proyectos (obras). Toda la información del sistema se organiza en esta jerarquía: '
    'Empresa → Obra → Informes / Remisiones.'
)

doc.add_heading('5.1 Lista de Empresas', level=2)
doc.add_paragraph(
    'Acceda desde el menú lateral: Directorio Empresas. Aquí verá todas las empresas '
    'constructoras registradas.'
)

doc.add_heading('Funcionalidades:', level=3)
items_empresas = [
    "Búsqueda: Escriba en el campo de búsqueda para filtrar por nombre, código o NIT.",
    "Filtro por ciudad: Use el selector de ciudad para ver solo empresas de una ubicación.",
    "Información visible: Nombre, código, NIT, ciudad, cantidad de obras, cantidad de usuarios.",
    "Clic en una empresa: Abre el expediente con todas sus obras.",
]
for item in items_empresas:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 Expediente de Empresa (Detalle)', level=2)
doc.add_paragraph(
    'Al hacer clic en una empresa, verá su expediente con la lista de todas sus obras '
    'ordenadas por fecha de creación (más recientes primero). Cada obra muestra el conteo '
    'de informes disponibles.'
)

doc.add_heading('5.3 Expediente de Obra (Detalle)', level=2)
doc.add_paragraph(
    'Al hacer clic en una obra, accede al expediente completo del proyecto. Esta es una de '
    'las pantallas más importantes del sistema.'
)

doc.add_heading('Secciones de la pantalla:', level=3)

doc.add_heading('Encabezado', level=4)
doc.add_paragraph(
    'Muestra el nombre de la obra, código, ciudad y fecha de creación. '
    'Incluye botones de acción rápida:'
)
items_header = [
    "Hoja de Trabajo: Abre la hoja de trabajo (ensayos) de esta obra.",
    "Nueva Remisión: Crea una nueva remisión de muestras para esta obra.",
    "Cargar Informe: Sube un nuevo informe PDF a esta obra.",
    "Lista de Precios ($): Gestiona los precios de servicios para facturación.",
    "Editar (lápiz): Modifica los datos de la obra.",
]
for item in items_header:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Panel Lateral Izquierdo', level=4)
items_lateral = [
    "Resumen: Conteo de documentos, remisiones y estado de la obra.",
    "Cliente: Datos de la empresa (nombre, NIT, ciudad) con enlace al expediente.",
    "Actividad: Línea de tiempo con fecha de creación, último informe y última remisión.",
]
for item in items_lateral:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Panel Principal Derecho', level=4)
items_derecho = [
    "Documentación Técnica: Tabla con los últimos 5 informes cargados. Muestra título, fecha y botón de descarga. Enlace para ver todos los informes.",
    "Remisiones de Muestras: Tabla con las últimas 5 remisiones. Muestra orden de trabajo, estado (Completada/Pendiente/Vencida), cantidad de muestras y fecha. Enlace para ver todas las remisiones.",
]
for item in items_derecho:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.4 Editar Empresa u Obra', level=2)
doc.add_paragraph(
    'Use el botón de editar (ícono de lápiz) para modificar los datos. Los campos editables son:'
)
add_simple_table(
    ["Entidad", "Campos Editables"],
    [
        ["Empresa", "Nombre, Código, NIT, Ciudad, Teléfono, Email"],
        ["Obra", "Nombre, Código de Obra, Constructora, Descripción, Localización"],
    ]
)

add_page_break()


# ============================================================
# 6. GESTIÓN DE INFORMES TÉCNICOS
# ============================================================
doc.add_heading('6. Gestión de Informes Técnicos', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'Los informes técnicos son documentos PDF generados por Geolab que contienen los resultados '
    'de los ensayos realizados. El sistema permite cargarlos, organizarlos por obra y distribuirlos '
    'a los clientes de forma segura.'
)

doc.add_heading('6.1 Cargar un Informe (Staff)', level=2)
add_colored_paragraph("Acceso: Menú lateral → Cargar Informe", GRAY_TEXT, size=10)
doc.add_paragraph()

add_step(1, "Seleccionar la empresa", "Elija la constructora del listado desplegable.")
add_step(2, "Seleccionar la obra", "El sistema carga automáticamente las obras de esa empresa.")
add_step(3, "Escribir el título del informe", "Un nombre descriptivo para identificar el documento.")
add_step(4, "Seleccionar el archivo PDF", "Haga clic en el botón de selección y elija el archivo desde su computador.")
add_step(5, "Hacer clic en 'Cargar Informe'", "El sistema sube el archivo y lo asocia a la obra seleccionada.")

doc.add_paragraph()
add_tip_box("También puede", "cargar informes directamente desde el expediente de una obra, usando el botón 'Cargar Informe' en el encabezado. Esto pre-selecciona la obra.")
add_warning_box("Solo se permiten archivos PDF. El sistema valida el formato tanto en el navegador como en el servidor.")

doc.add_heading('6.2 Consultar Informes por Obra', level=2)

doc.add_paragraph(
    'Desde el expediente de cualquier obra, haga clic en "Ver todos" en la sección de '
    'Documentación Técnica para acceder a la lista completa de informes.'
)

doc.add_heading('Herramientas de búsqueda:', level=3)
items_busqueda = [
    "Búsqueda por título: Escriba parte del nombre del informe.",
    "Filtro por fecha: Seleccione un rango de fechas (desde - hasta).",
    "Paginación: Los informes se muestran de 20 en 20.",
]
for item in items_busqueda:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3 Descargar Informes (Clientes)', level=2)
doc.add_paragraph(
    'Los clientes (Director y Residente) acceden a los informes desde su panel:'
)
add_step(1, "Iniciar sesión", "El sistema muestra sus obras disponibles.")
add_step(2, "Seleccionar una obra", "Haga clic en la tarjeta de la obra.")
add_step(3, "Buscar el informe", "Use la búsqueda por título o los filtros de fecha.")
add_step(4, "Descargar", "Haga clic en el ícono de descarga para obtener el PDF.")

add_page_break()


# ============================================================
# 7. REMISIONES DE MUESTRAS
# ============================================================
doc.add_heading('7. Remisiones de Muestras (F-GC-05)', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'La remisión de muestras es el documento que registra el envío de muestras de campo al '
    'laboratorio. En el sistema se maneja de forma digital reemplazando el formato físico '
    'F-GC-05. Cada remisión incluye una orden de trabajo, los datos del remitente y el detalle '
    'de cada muestra enviada.'
)

doc.add_heading('7.1 Flujo General de una Remisión', level=2)

doc.add_paragraph(
    'El proceso de una remisión sigue estos pasos:'
)

add_step(1, "Creación", "El remitente o el staff crea la remisión con los datos de las muestras.")
add_step(2, "Recepción automática", "Al enviarse, el sistema crea automáticamente la Hoja de Trabajo y los registros de resultado para cada muestra.")
add_step(3, "Notificación", "El staff recibe una notificación (campana + email) de la nueva remisión.")
add_step(4, "Ensayos", "Los técnicos de laboratorio registran los resultados de cada muestra en la fecha programada.")
add_step(5, "Informes", "Se generan los informes técnicos y se cargan al sistema para el cliente.")

doc.add_heading('7.2 Crear una Remisión (Remitente)', level=2)
add_colored_paragraph("Acceso: Panel de Remitente → Seleccionar obra → Nueva Remisión", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El remitente es quien está en campo y envía las muestras al laboratorio. El proceso es:'
)

add_step(1, "Ir al Panel de Remitente", "Al iniciar sesión como remitente, verá sus obras asignadas.")
add_step(2, "Seleccionar la obra", "Haga clic en el botón de nueva remisión junto a la obra.")
add_step(3, "Llenar datos generales", "Orden de trabajo, observaciones generales de la remisión.")
add_step(4, "Agregar muestras", "Para cada muestra, ingrese:")
doc.add_paragraph()

muestras_campos = [
    "Número de muestra (identificador secuencial)",
    "Tipo de muestra (cilindro, vigueta, cubo, núcleo, etc.)",
    "Ensayos a realizar",
    "Localización de donde se tomó la muestra",
    "Fecha de toma de la muestra",
    "Edad de ensayo en días (ej: 7, 14, 28 días)",
    "F'c o resistencia de diseño (en MPa)",
]
for item in muestras_campos:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(2)

add_step(5, "Enviar la remisión", "Al hacer clic en 'Enviar', el sistema registra la remisión con sus datos de firma (nombre, cargo, fecha, hora, IP).")

doc.add_paragraph()
add_tip_box("Firma automática", "El sistema registra automáticamente su nombre, cargo, email, fecha, hora e IP como firma digital de la remisión.")
add_warning_box("Una vez enviada, la remisión no se puede modificar. Verifique todos los datos antes de enviar.")

doc.add_heading('7.3 Crear una Remisión (Staff)', level=2)
add_colored_paragraph("Acceso: Expediente de Obra → Nueva Remisión", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El staff también puede crear remisiones desde el expediente de cualquier obra. '
    'El proceso es similar al del remitente pero con la ventaja de poder seleccionar '
    'cualquier obra del sistema.'
)

doc.add_heading('7.4 Consultar Remisiones', level=2)

doc.add_heading('Vista del Staff:', level=3)
doc.add_paragraph(
    'Acceda desde el menú lateral → Remisiones de Muestras. Verá todas las remisiones '
    'del sistema con:'
)
items_remision_staff = [
    "Búsqueda por orden de trabajo, obra o constructora",
    "Filtro por estado (enviada, completada)",
    "Enlace al detalle de cada remisión",
    "Botón para descargar el PDF de la remisión",
]
for item in items_remision_staff:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Vista del Remitente:', level=3)
doc.add_paragraph(
    'Acceda desde el menú lateral → Mis Remisiones. Verá solo sus remisiones con:'
)
items_remision_rem = [
    "Filtro por obra",
    "Búsqueda por orden de trabajo",
    "Filtro por estado",
    "Vista de detalle en solo lectura",
]
for item in items_remision_rem:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.5 Detalle de una Remisión', level=2)
doc.add_paragraph(
    'Al abrir el detalle de una remisión verá:'
)
items_detalle_rem = [
    "Datos generales: obra, constructora, orden de trabajo, estado, fechas",
    "Datos del firmante: nombre, cargo, email, fecha y hora de firma",
    "Tabla de muestras: número, tipo, ensayos, localización, fecha de toma, edad, resistencia",
    "Botón para descargar el PDF de la remisión en formato F-GC-05",
]
for item in items_detalle_rem:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.6 Estados de una Remisión', level=2)

add_simple_table(
    ["Estado", "Significado", "Color"],
    [
        ["Completada", "La remisión fue recibida y procesada exitosamente", "Verde"],
        ["Pendiente", "La remisión fue enviada pero aún no ha sido completada", "Amarillo"],
        ["Vencida", "La remisión excedió el tiempo de respuesta esperado", "Rojo"],
    ]
)

doc.add_heading('7.7 PDF de Remisión', level=2)
doc.add_paragraph(
    'Cada remisión genera automáticamente un PDF en formato F-GC-05 que incluye todos los '
    'datos de la remisión, la tabla de muestras y la información de cadena de custodia. '
    'Este PDF se puede descargar desde el detalle de la remisión o desde la lista.'
)

add_page_break()


# ============================================================
# 8. ENSAYOS DE LABORATORIO
# ============================================================
doc.add_heading('8. Ensayos de Laboratorio (F-GT-05)', level=1)
doc.add_paragraph()
add_colored_paragraph("Acceso: Staff Administrativo y Técnico de Laboratorio", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El módulo de ensayos permite a los técnicos de laboratorio registrar los resultados '
    'de las pruebas realizadas sobre las muestras recibidas. El sistema se basa en el formato '
    'F-GT-05 (Hoja de Trabajo) y calcula automáticamente los indicadores clave de cada ensayo.'
)

doc.add_heading('8.1 Conceptos Clave', level=2)

conceptos = [
    ("Hoja de Trabajo", "Documento digital que agrupa todos los resultados de una remisión. Se crea automáticamente al recibir una remisión."),
    ("Resultado de Muestra", "Registro individual de las mediciones y cálculos de cada muestra."),
    ("Fecha de Falla", "Fecha programada para el ensayo = Fecha de toma + Edad de ensayo en días."),
    ("Esfuerzo (MPa)", "Resultado principal del ensayo: resistencia a la compresión de la muestra."),
    ("Porcentaje de Desarrollo", "Relación entre el esfuerzo obtenido y la resistencia de diseño (F'c)."),
]
for titulo, desc in conceptos:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run2 = p.add_run(desc)

doc.add_heading('8.2 Ensayos del Día', level=2)
add_colored_paragraph("Acceso: Menú lateral → Ensayos del Día", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Esta pantalla muestra todas las muestras que deben ser ensayadas hoy. Es la vista '
    'principal para el trabajo diario del técnico.'
)

doc.add_heading('Secciones:', level=3)
secciones_hoy = [
    ("Muestras para HOY (rojo)", "Muestras cuya fecha de falla es exactamente hoy. Estas son la prioridad del día."),
    ("Muestras Vencidas", "Muestras cuya fecha de falla ya pasó y aún no han sido ensayadas. Requieren atención inmediata."),
    ("Próximas (3 días)", "Muestras que serán ensayadas en los próximos 3 días, para planificación."),
]
for titulo, desc in secciones_hoy:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run2 = p.add_run(desc)

doc.add_heading('8.3 Hojas de Trabajo', level=2)
add_colored_paragraph("Acceso: Menú lateral → Hojas de Trabajo", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Muestra la lista de obras que tienen hojas de trabajo activas. Para cada obra se ve:'
)
items_hojas = [
    "Total de muestras de la obra",
    "Muestras completadas vs. pendientes",
    "Muestras programadas para hoy y vencidas",
    "Barra de progreso visual",
]
for item in items_hojas:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Al hacer clic en una obra, accede a la hoja de trabajo completa con todas las muestras '
    'de todas las remisiones de esa obra.'
)

doc.add_heading('8.4 Registrar Resultados de Ensayo', level=2)
doc.add_paragraph(
    'Dentro de la hoja de trabajo de una obra, el técnico puede editar los resultados de '
    'cada muestra directamente en la tabla (edición en línea).'
)

doc.add_heading('Campos a registrar:', level=3)

add_simple_table(
    ["Campo", "Descripción", "Ejemplo"],
    [
        ["Diámetro D1, D2, D3", "Tres mediciones del diámetro del cilindro (mm)", "152.1, 152.3, 151.9"],
        ["Longitud L1, L2, L3", "Tres mediciones de la longitud del cilindro (mm)", "305.2, 304.8, 305.0"],
        ["Peso (gramos)", "Peso de la muestra en gramos", "12,450"],
        ["Carga Máxima (kN)", "Carga máxima aplicada en la prensa", "520.3"],
        ["Forma de Falla", "Tipo de fractura observada (Tipo 1-6)", "Tipo 2"],
    ]
)

doc.add_heading('8.5 Cálculos Automáticos', level=2)
doc.add_paragraph(
    'Al registrar las mediciones, el sistema calcula automáticamente:'
)

calculos = [
    ("Diámetro Real", "Promedio de D1, D2 y D3"),
    ("Longitud Real", "Promedio de L1, L2 y L3"),
    ("Área (mm²)", "π × diámetro² / 4 / 100"),
    ("Esfuerzo (MPa)", "(Carga kN × 101.9716 / Área) / 10"),
    ("% de Desarrollo", "(Esfuerzo / F'c) × 100"),
    ("Cumple", "Si el porcentaje de desarrollo alcanza el mínimo esperado para la edad"),
]
for titulo, formula in calculos:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo} = ")
    run.bold = True
    run2 = p.add_run(formula)

doc.add_paragraph()
add_tip_box("Automatización", "No necesita hacer estos cálculos manualmente. Solo ingrese las mediciones y el sistema genera los resultados al instante.")

doc.add_heading('8.6 Estados de un Resultado', level=2)
add_simple_table(
    ["Estado", "Significado"],
    [
        ["Pendiente", "La muestra aún no ha sido ensayada"],
        ["Completado", "Todas las mediciones fueron registradas y los cálculos generados"],
    ]
)

add_page_break()


# ============================================================
# 9. SISTEMA DE FACTURACIÓN
# ============================================================
doc.add_heading('9. Sistema de Facturación', level=1)
doc.add_paragraph()
add_colored_paragraph("Acceso: Solo Staff Administrativo", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El módulo de facturación permite registrar los servicios prestados a cada obra, '
    'gestionar precios, generar facturas automáticas y llevar un histórico completo de '
    'la facturación. El menú lateral agrupa todas las funciones bajo la sección FACTURACIÓN.'
)

doc.add_heading('9.1 Panel de Facturación', level=2)
add_colored_paragraph("Acceso: Menú lateral → Panel Facturación", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El panel muestra los indicadores clave de facturación:'
)

kpis_fact = [
    ("Facturas del mes", "Cantidad y monto total de facturas generadas en el mes actual."),
    ("Registros del mes", "Servicios registrados en el mes actual."),
    ("Registros pendientes", "Servicios registrados que aún no han sido incluidos en una factura."),
    ("Top Clientes", "Los 3 clientes con mayor facturación histórica."),
    ("Top Servicios", "Los 3 servicios más realizados en los últimos 90 días."),
    ("Gráfico mensual", "Gráfico de barras con la facturación de los últimos 6 meses."),
]
for titulo, desc in kpis_fact:
    p = doc.add_paragraph()
    run = p.add_run(f"• {titulo}: ")
    run.bold = True
    run2 = p.add_run(desc)

doc.add_heading('9.2 Catálogo de Servicios', level=2)
add_colored_paragraph("Acceso: Menú lateral → Catálogo Servicios", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El catálogo organiza todos los servicios que ofrece Geolab en categorías. '
    'Cada servicio tiene un código único, nombre descriptivo y norma técnica asociada.'
)

doc.add_heading('Estructura del Catálogo', level=3)
doc.add_paragraph(
    'Los servicios se organizan en categorías. Ejemplo:'
)

add_simple_table(
    ["Categoría", "Código", "Ejemplos de Servicios"],
    [
        ["1 - CONCRETOS", "I-1, I-2...", "Resistencia a la compresión, flexión, etc."],
        ["2 - SUELOS", "II-1, II-2...", "Granulometría, límites de Atterberg, etc."],
        ["3 - ASFALTOS", "III-1, III-2...", "Estabilidad Marshall, contenido de asfalto, etc."],
        ["7 - TRANSPORTE", "VII-1...", "Transporte de muestras (regla IVA especial)"],
    ]
)

doc.add_paragraph()
add_warning_box(
    "La categoría TRANSPORTE (código 7) tiene una regla especial de IVA: el IVA solo se "
    "aplica si se selecciona explícitamente al generar la factura. Todos los demás servicios "
    "siempre llevan IVA."
)

doc.add_heading('Gestionar el Catálogo', level=3)
items_catalogo = [
    "Crear categoría: Botón 'Nueva Categoría', ingrese código y nombre.",
    "Crear servicio: Botón 'Nuevo Servicio', seleccione categoría e ingrese código, nombre y norma.",
    "Editar: Use el ícono de lápiz junto a cada categoría o servicio.",
    "Eliminar servicio: Use el ícono de papelera (solo si no tiene registros asociados).",
    "Buscar: Use el buscador para filtrar por código, nombre o norma.",
]
for item in items_catalogo:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.3 Lista de Precios por Obra', level=2)
add_colored_paragraph("Acceso: Expediente de Obra → Botón $ (Lista de Precios)", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Cada obra puede tener precios diferentes para cada servicio. Antes de poder registrar '
    'servicios para una obra, debe configurar su lista de precios.'
)

add_step(1, "Abrir el expediente de la obra", "Navegue al detalle de la obra.")
add_step(2, "Hacer clic en el botón '$'", "Se abre la lista de precios organizada por categoría de servicio.")
add_step(3, "Ingresar los precios", "Para cada servicio que aplique a esta obra, ingrese el precio unitario.")
add_step(4, "Guardar", "Haga clic en 'Guardar Precios'. Solo se guardan los servicios con precio ingresado.")

doc.add_paragraph()
add_tip_box("Nota", "Los servicios sin precio asignado no estarán disponibles al registrar servicios para esa obra.")

doc.add_heading('9.4 Impuestos', level=2)
add_colored_paragraph("Acceso: Menú lateral → Impuestos", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Gestione los impuestos que se aplican a las facturas. El impuesto principal es el IVA.'
)

add_step(1, "Crear el IVA", "Haga clic en 'Nuevo Impuesto', ingrese nombre: IVA, porcentaje: 19.00, estado: Activo.")
add_step(2, "Editar si cambia", "Use el ícono de lápiz para modificar el porcentaje si es necesario.")

doc.add_paragraph()
add_tip_box("Primera vez", "Al instalar el sistema, debe crear el impuesto IVA (19%) desde esta pantalla antes de generar facturas.")

doc.add_heading('9.5 Registrar Servicios', level=2)
add_colored_paragraph("Acceso: Menú lateral → Registrar Servicio", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Aquí se registran los servicios que Geolab ha realizado para una obra. Estos registros '
    'son la base para generar las facturas.'
)

add_step(1, "Seleccionar constructora", "Busque y seleccione la empresa constructora.")
add_step(2, "Seleccionar obra", "El sistema carga las obras de esa empresa. Seleccione una.")
add_step(3, "Indicar la fecha", "Seleccione la fecha en que se realizó el servicio.")
add_step(4, "Agregar líneas de servicio", "Para cada servicio realizado:")

items_linea = [
    "Busque el servicio: El buscador filtra por código o nombre.",
    "Ingrese el número de informe: Referencia del informe asociado al servicio.",
    "Indique la cantidad: Número de ensayos o servicios realizados.",
    "El sistema muestra automáticamente el precio unitario (de la lista de precios de la obra) y calcula el subtotal.",
]
for item in items_linea:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(2)

add_step(5, "Agregar más líneas si es necesario", "Use el botón '+ Agregar línea' para registrar más servicios en la misma fecha.")
add_step(6, "Guardar", "Haga clic en 'Guardar Registros'. El sistema guarda todos los servicios de forma atómica.")

doc.add_paragraph()
add_tip_box("Congelamiento de precio", "Al guardar, el precio se 'congela' en el registro. Si después cambia el precio en la lista, los registros ya guardados no se afectan.")
add_warning_box("Si un servicio no tiene precio asignado en la lista de precios de la obra, no se podrá registrar.")

doc.add_heading('9.6 Histórico de Registros', level=2)
add_colored_paragraph("Acceso: Menú lateral → Histórico", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Muestra todos los registros de servicios con opciones de filtrado avanzado.'
)

doc.add_heading('Filtros disponibles:', level=3)
filtros_hist = [
    "Constructora: Busque por nombre de empresa.",
    "Obra: Filtre por obra específica (se carga al seleccionar constructora).",
    "Servicio: Busque por tipo de servicio.",
    "Fecha desde / hasta: Rango de fechas.",
]
for item in filtros_hist:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Información de cada registro:', level=3)
doc.add_paragraph(
    'La tabla muestra: fecha, constructora, obra, servicio, número de informe, cantidad, '
    'precio unitario, subtotal y estado (Pendiente o Facturado #ID).'
)

doc.add_heading('Acciones disponibles:', level=3)
items_acciones_hist = [
    "Editar: Solo registros no facturados. Permite modificar los datos del servicio.",
    "Eliminar: Solo registros no facturados. Elimina el registro con confirmación.",
]
for item in items_acciones_hist:
    doc.add_paragraph(item, style='List Bullet')

add_warning_box("Los registros que ya están en una factura no se pueden editar ni eliminar. Para liberarlos, primero anule la factura.")

doc.add_heading('9.7 Generar Factura', level=2)
add_colored_paragraph("Acceso: Menú lateral → Generar Factura", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Este es el proceso para generar una factura a partir de los registros pendientes de una obra.'
)

add_step(1, "Seleccionar constructora y obra", "Busque y seleccione la empresa y obra a facturar.")
add_step(2, "Definir el periodo", "Seleccione la fecha de inicio y fin del periodo a facturar.")
add_step(3, "Buscar registros", "Haga clic en 'Buscar registros'. El sistema muestra todos los registros pendientes del periodo.")

doc.add_paragraph()
doc.add_paragraph(
    'La vista previa muestra dos secciones:'
)

doc.add_heading('Servicios Normales:', level=4)
doc.add_paragraph(
    'Tabla con todos los servicios que no son de transporte. El IVA se aplica automáticamente '
    'sobre el total de estos servicios.'
)

doc.add_heading('Servicios de Transporte:', level=4)
doc.add_paragraph(
    'Tabla separada para los servicios de transporte (categoría 7). Cada servicio de transporte '
    'tiene un checkbox para indicar si lleva IVA o no. Al marcar/desmarcar, el total se '
    'recalcula en tiempo real.'
)

doc.add_heading('Resumen:', level=4)
doc.add_paragraph(
    'En la parte inferior se muestra: subtotal general, monto de IVA calculado y total de la factura.'
)

add_step(4, "Confirmar y generar", "Haga clic en 'Generar Factura'. El sistema pide confirmación antes de proceder.")

doc.add_paragraph()
doc.add_paragraph(
    'Al generar la factura, el sistema automáticamente:'
)
items_generar = [
    "Calcula el IVA según las reglas (servicios normales + transporte seleccionado)",
    "Crea el registro de factura con todos los montos",
    "Vincula los registros de servicio a la factura (ya no aparecen como pendientes)",
    "Genera un PDF de la factura con el formato profesional de Geolab",
]
for item in items_generar:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.8 Repositorio de Facturas', level=2)
add_colored_paragraph("Acceso: Menú lateral → Repositorio", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El repositorio organiza todas las facturas en una navegación de tres niveles:'
)

add_step(1, "Nivel 1 - Clientes", "Ve todas las empresas con facturas, mostrando el total facturado a cada una.")
add_step(2, "Nivel 2 - Obras del Cliente", "Al hacer clic en una empresa, ve sus obras con conteo de facturas.")
add_step(3, "Nivel 3 - Facturas de la Obra", "Al hacer clic en una obra, ve la tabla detallada de facturas.")

doc.add_heading('Acciones sobre facturas:', level=3)

add_simple_table(
    ["Acción", "Descripción", "Condición"],
    [
        ["Ver PDF", "Descarga el PDF de la factura", "Solo si tiene PDF generado"],
        ["Marcar Pagada", "Cambia el estado a PAGADA", "Solo facturas PENDIENTES"],
        ["Anular", "Anula la factura y libera los registros", "Solo facturas PENDIENTES"],
    ]
)

doc.add_heading('9.9 Anulación de Facturas', level=2)
doc.add_paragraph(
    'Anular una factura es un proceso reversible que:'
)
items_anular = [
    "Marca la factura como ANULADA (no se elimina del sistema)",
    "Desvincula los registros de servicio de la factura",
    "Los registros quedan disponibles nuevamente para ser incluidos en una nueva factura",
    "La factura anulada se muestra tachada en el repositorio",
]
for item in items_anular:
    doc.add_paragraph(item, style='List Bullet')

add_warning_box("Solo se pueden anular facturas en estado PENDIENTE. Las facturas ya marcadas como PAGADA no se pueden anular.")

doc.add_heading('9.10 Estados de una Factura', level=2)

add_simple_table(
    ["Estado", "Significado", "Color", "Acciones Permitidas"],
    [
        ["PENDIENTE", "Factura generada, pendiente de pago", "Amarillo", "Marcar pagada, Anular, Ver PDF"],
        ["PAGADA", "Factura cobrada y pagada por el cliente", "Verde", "Ver PDF"],
        ["ANULADA", "Factura cancelada, registros liberados", "Rojo", "Ver PDF (histórico)"],
    ]
)

add_page_break()


# ============================================================
# 10. SISTEMA DE CALIDAD
# ============================================================
doc.add_heading('10. Sistema de Calidad', level=1)
doc.add_paragraph()
add_colored_paragraph("Acceso: Staff Administrativo y Técnicos (según permisos)", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El módulo de Calidad permite gestionar toda la documentación del Sistema de Gestión '
    'de Calidad de Geolab. Los documentos se organizan en áreas temáticas con una estructura '
    'jerárquica de carpetas.'
)

doc.add_heading('10.1 Explorador de Calidad', level=2)
add_colored_paragraph("Acceso: Menú lateral → Sistema de Calidad", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'La pantalla principal muestra las áreas de calidad en formato de cuadrícula. '
    'Cada área muestra el nombre, la cantidad de carpetas y documentos que contiene.'
)

doc.add_paragraph(
    'Los administradores ven todas las áreas. Los demás usuarios de Geolab solo ven '
    'las áreas a las que tienen acceso asignado.'
)

doc.add_heading('10.2 Navegación de Carpetas', level=2)
doc.add_paragraph(
    'Al hacer clic en un área, se abre la vista de carpetas con navegación tipo explorador '
    'de archivos:'
)
items_carpetas = [
    "Breadcrumb (ruta de navegación): Muestra la ruta actual (Área → Carpeta → Subcarpeta).",
    "Carpetas: Se muestran como íconos con nombre. Haga clic para entrar.",
    "Documentos: Se muestran con nombre, tamaño y fecha. Haga clic para descargar.",
]
for item in items_carpetas:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.3 Subir Documentos', level=2)
doc.add_paragraph(
    'Dentro de una carpeta (si tiene permiso de edición):'
)
add_step(1, "Navegar a la carpeta destino", "Abra la carpeta donde desea subir el archivo.")
add_step(2, "Hacer clic en 'Subir Documento'", "Se abre el formulario de carga.")
add_step(3, "Seleccionar el archivo", "Puede seleccionar múltiples archivos a la vez.")
add_step(4, "Confirmar", "El sistema sube los archivos y registra quién y cuándo los subió.")

doc.add_heading('10.4 Crear Carpetas', level=2)
doc.add_paragraph(
    'Puede crear nuevas carpetas dentro de un área o dentro de otra carpeta. '
    'Use el botón "Nueva Carpeta" e ingrese el nombre.'
)

doc.add_heading('10.5 Eliminar', level=2)
add_colored_paragraph("Acceso: Solo Administradores", GRAY_TEXT, size=10)
doc.add_paragraph()
items_eliminar = [
    "Documentos: Se eliminan con confirmación previa.",
    "Carpetas: Solo se pueden eliminar si están vacías (sin subcarpetas ni documentos).",
]
for item in items_eliminar:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.6 Control de Acceso por Áreas', level=2)
add_colored_paragraph("Acceso: Menú Administración (solo Administradores)", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'El administrador puede controlar qué usuarios tienen acceso a cada área y con qué nivel:'
)

add_simple_table(
    ["Nivel de Acceso", "Puede ver", "Puede subir/crear", "Puede eliminar"],
    [
        ["Sin acceso", "No", "No", "No"],
        ["Lectura", "Sí", "No", "No"],
        ["Edición", "Sí", "Sí", "No (solo admin)"],
    ]
)

doc.add_paragraph()
add_tip_box("Administradores", "Los usuarios administradores tienen acceso completo a todas las áreas automáticamente, sin necesidad de asignación individual.")

add_page_break()


# ============================================================
# 11. GESTIÓN DE USUARIOS
# ============================================================
doc.add_heading('11. Gestión de Usuarios y Accesos', level=1)
doc.add_paragraph()
add_colored_paragraph("Acceso: Solo Staff Administrativo", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Desde aquí se crean, editan y administran los usuarios que pueden acceder al portal. '
    'Esto incluye tanto personal de Geolab como usuarios externos (clientes).'
)

doc.add_heading('11.1 Lista de Usuarios', level=2)
add_colored_paragraph("Acceso: Menú lateral → Usuarios y Accesos", GRAY_TEXT, size=10)
doc.add_paragraph()

doc.add_paragraph(
    'Muestra todos los usuarios del sistema con herramientas de búsqueda y filtrado:'
)
items_lista_usuarios = [
    "Búsqueda: Por nombre, email, usuario o empresa.",
    "Filtro por rol: Director, Residente o Remitente.",
    "Filtro por empresa: Ver solo usuarios de una empresa específica.",
    "Estadísticas: Total de directores y residentes en el sistema.",
    "Estado: Indicador visual de activo/inactivo con interruptor rápido.",
]
for item in items_lista_usuarios:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('11.2 Crear un Nuevo Usuario', level=2)
add_colored_paragraph("Botón: + Nuevo Usuario", GRAY_TEXT, size=10)
doc.add_paragraph()

add_step(1, "Datos personales", "Nombre, apellido, email y nombre de usuario.")
add_step(2, "Contraseña", "Defina una contraseña segura para el usuario.")
add_step(3, "Perfil del usuario", "")

items_perfil = [
    "Empresa: Seleccione a qué constructora pertenece.",
    "Rol: Director (ve todas las obras), Residente (solo obras asignadas), o Remitente (crea remisiones).",
    "Cargo: Cargo del usuario en la empresa.",
    "Teléfono: Número de contacto.",
]
for item in items_perfil:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(2)

add_step(4, "Obras asignadas (solo Residente/Remitente)", "Si el rol es Residente o Remitente, seleccione las obras a las que tendrá acceso.")
add_step(5, "Guardar", "El usuario queda activo inmediatamente y puede iniciar sesión.")

doc.add_heading('11.3 Editar un Usuario', level=2)
doc.add_paragraph(
    'Haga clic en el nombre del usuario o en el botón de editar. Puede modificar todos '
    'los datos excepto el nombre de usuario. Si necesita cambiar la contraseña, use el '
    'campo de contraseña del formulario.'
)

doc.add_heading('11.4 Activar / Desactivar Usuarios', level=2)
doc.add_paragraph(
    'Use el interruptor de estado en la lista para activar o desactivar un usuario rápidamente. '
    'Un usuario desactivado no puede iniciar sesión pero sus datos se conservan en el sistema.'
)

add_tip_box("Recomendación", "En lugar de eliminar un usuario, desactívelo. Así conserva el historial de sus acciones en el sistema.")

add_page_break()


# ============================================================
# 12. NOTIFICACIONES
# ============================================================
doc.add_heading('12. Notificaciones', level=1)
doc.add_paragraph()

doc.add_paragraph(
    'El sistema incluye un mecanismo de notificaciones en tiempo real para mantener informados '
    'a los usuarios sobre eventos importantes.'
)

doc.add_heading('12.1 Campana de Notificaciones', level=2)
doc.add_paragraph(
    'En la barra superior derecha encontrará el ícono de campana. Un número en rojo indica '
    'la cantidad de notificaciones sin leer. Haga clic para ver las más recientes.'
)

doc.add_heading('12.2 Tipos de Notificaciones', level=2)

add_simple_table(
    ["Evento", "¿Quién la recibe?", "Contenido"],
    [
        ["Nueva remisión creada", "Staff administrativo", "Datos de la remisión, obra y remitente"],
        ["Remisión completada", "Staff administrativo", "Confirmación de recepción en laboratorio"],
    ]
)

doc.add_heading('12.3 Gestionar Notificaciones', level=2)
add_colored_paragraph("Acceso: Clic en la campana → Ver todas", GRAY_TEXT, size=10)
doc.add_paragraph()

items_notif = [
    "Ver todas: Lista completa de notificaciones con paginación.",
    "Marcar como leída: Haga clic en una notificación individual.",
    "Marcar todas como leídas: Botón para limpiar el contador de una sola vez.",
    "Enlace directo: Cada notificación enlaza al detalle del evento relacionado.",
]
for item in items_notif:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_tip_box("Actualización automática", "El contador de notificaciones se actualiza automáticamente cada 30 segundos sin necesidad de recargar la página.")

doc.add_heading('12.4 Notificaciones por Email', level=2)
doc.add_paragraph(
    'Además de la campana en el sistema, las notificaciones más importantes (como nuevas '
    'remisiones) también se envían por correo electrónico a los usuarios correspondientes.'
)

add_page_break()


# ============================================================
# 13. PREGUNTAS FRECUENTES
# ============================================================
doc.add_heading('13. Preguntas Frecuentes', level=1)
doc.add_paragraph()

faqs = [
    ("¿Cómo recupero mi contraseña?",
     "Contacte al administrador de Geolab para que restablezca su contraseña. "
     "No existe opción de recuperación automática por seguridad."),

    ("¿Por qué no veo todas las obras de mi empresa?",
     "Si su rol es Residente, solo ve las obras donde fue asignado. Solicite al "
     "administrador que lo asigne a las obras adicionales que necesita. Si es Director, "
     "debería ver todas las obras."),

    ("¿Puedo modificar una remisión después de enviarla?",
     "No. Una vez enviada, la remisión queda registrada con firma digital y no puede "
     "modificarse. Si hay un error, contacte al staff de Geolab."),

    ("¿Cómo agrego un nuevo servicio al catálogo de facturación?",
     "Vaya a Catálogo Servicios → Nuevo Servicio. Asigne una categoría, código, nombre "
     "y norma técnica."),

    ("¿Puedo facturar servicios de diferentes obras en una sola factura?",
     "No. Cada factura corresponde a una sola obra y un periodo de tiempo específico."),

    ("¿Qué pasa cuando anulo una factura?",
     "La factura se marca como ANULADA y los registros de servicio quedan libres para "
     "incluirlos en una nueva factura. La factura anulada permanece visible en el "
     "repositorio como registro histórico."),

    ("¿Por qué un servicio me aparece 'Sin precio' al registrarlo?",
     "Porque no se ha configurado el precio de ese servicio para esa obra. Vaya al "
     "expediente de la obra → botón $ → configure el precio del servicio."),

    ("¿Cómo sé qué muestras debo ensayar hoy?",
     "El sistema calcula automáticamente la fecha de falla (fecha de toma + días de edad). "
     "Vaya a Ensayos del Día para ver las muestras programadas para hoy."),

    ("¿Los cálculos de los ensayos se hacen automáticamente?",
     "Sí. Solo necesita ingresar las mediciones (diámetros, longitudes, peso, carga y "
     "forma de falla). El sistema calcula esfuerzo, porcentaje de desarrollo y cumplimiento."),

    ("¿Puedo acceder desde el celular?",
     "Sí. El portal es responsive y se adapta a cualquier dispositivo. Aunque para tareas "
     "de laboratorio y facturación se recomienda usar computador por comodidad."),

    ("¿Quién puede ver los documentos de Calidad?",
     "Solo el personal de Geolab con acceso asignado. Los administradores ven todo. "
     "Los demás usuarios solo ven las áreas asignadas por el administrador."),

    ("¿Cómo sé si una muestra cumplió con la resistencia esperada?",
     "El sistema calcula el porcentaje de desarrollo (esfuerzo obtenido vs. resistencia "
     "de diseño F'c) y marca automáticamente si cumple o no."),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    run = p.add_run(pregunta)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(respuesta)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(12)

add_page_break()

# ============================================================
# CONTRAPORTADA
# ============================================================
for _ in range(8):
    doc.add_paragraph()

add_colored_paragraph("GEOLAB S.A.S", BRAND_BLUE, bold=True, size=24,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
add_colored_paragraph("Portal de Gestión Técnica", DARK_BLUE, size=16,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_colored_paragraph("Manual de Usuario v1.0", GRAY_TEXT, size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
add_colored_paragraph("Marzo 2026", GRAY_TEXT, size=11,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

add_colored_paragraph(
    "Para soporte técnico contacte al equipo de Geolab S.A.S",
    GRAY_TEXT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER
)
add_colored_paragraph(
    "https://portalgeolab.com",
    BRAND_BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_paragraph()
doc.add_paragraph()
add_colored_paragraph(
    "Desarrollado por Vadom Data Consulting SAS",
    GRAY_TEXT, size=9, align=WD_ALIGN_PARAGRAPH.CENTER
)


# ============================================================
# GUARDAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Manual_Usuario_Geolab.docx')
doc.save(output_path)
print(f"\n✅ Manual generado exitosamente: {output_path}")
print(f"   Tamaño: {os.path.getsize(output_path) / 1024:.0f} KB")
